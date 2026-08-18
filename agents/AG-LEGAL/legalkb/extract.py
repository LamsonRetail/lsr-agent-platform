"""Trích text từ file người dùng gửi (PDF/DOCX/TXT) — cho S3 và S5.

Làm TRONG agent, không dùng `/v1/extract` của platform: endpoint đó yêu cầu quyền admin
(`platform_api/app.py`, `_require_admin`) nên token agent gọi sẽ 403.
"""
import io
import re
import zipfile

MAX_CHARS = 120_000        # đủ cho hợp đồng dài; chặn file rác làm nổ prompt


class ExtractError(RuntimeError):
    pass


def from_bytes(data, file_name=""):
    """Trả text đã chuẩn hoá khoảng trắng. Không đoán định dạng theo nội dung — theo
    phần mở rộng, vì đoán sai sẽ ra text rác mà vẫn 'thành công'."""
    name = (file_name or "").lower()
    if name.endswith(".pdf"):
        text = _pdf(data)
    elif name.endswith(".docx"):
        text = _docx(data)
    elif name.endswith((".txt", ".md")):
        text = data.decode("utf-8", "replace")
    else:
        raise ExtractError(
            f"chưa đọc được định dạng của '{file_name or 'file'}'. "
            f"Gửi lại dạng PDF hoặc DOCX giúp mình.")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        raise ExtractError("file không có nội dung văn bản đọc được "
                           "(có thể là bản scan ảnh — cần bản gốc có text).")
    return text[:MAX_CHARS]


def _pdf(data):
    from pypdf import PdfReader
    try:
        r = PdfReader(io.BytesIO(data))
        return "\n".join((p.extract_text() or "") for p in r.pages)
    except Exception as exc:
        raise ExtractError(f"không đọc được PDF: {exc}")


def _docx(data):
    from docx import Document
    try:
        doc = Document(io.BytesIO(data))
    except (zipfile.BadZipFile, KeyError) as exc:
        raise ExtractError(f"không đọc được DOCX: {exc}")
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:                       # điều khoản hợp đồng hay nằm trong bảng
        for row in t.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(parts)
