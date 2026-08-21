"""S2 — tạo hợp đồng từ template (PLAN Phase 3).

Luồng: chọn template → hỏi lần lượt field còn thiếu → xác nhận → điền docx → upload
Drive → **gate Pháp chế** → chỉ khi được duyệt mới gửi link cho người yêu cầu.

State đa lượt nằm ở bảng `contract_drafts`, KHÔNG nhồi vào prompt: người dùng bỏ dở hôm
nay, mai quay lại vẫn tiếp đúng field đang hỏi, kể cả sau khi restart container.

Registry template do legal team quản trên Drive: mỗi hợp đồng gồm
  - `<tên>.docx`        chứa placeholder {{ten_ben_a}}, {{gia_tri}}, …
  - `<tên>.fields.json` mô tả field  [{"key","label","required","hint"}]
Không có file .fields.json thì tự dò placeholder trong docx (mọi field coi là bắt buộc).
"""
import io
import json
import re
import time
import zipfile

DRAFT_MARK = "DRAFT — CHƯA CÓ HIỆU LỰC PHÁP LÝ"
PLACEHOLDER = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")

# Mẫu THẬT của legal team không dùng `{{...}}` — kiểm mẫu "Hop dong mua ban_Hapas"
# ngày 21/08: chỗ cần điền là **dãy dấu ba chấm chèn giữa câu**:
#   "Số: ………/2026/HDMBHH/…….- HTC"
#   "Thời gian giao hàng: …….. giờ ngày …………….;"
# Nên phải nhận cả kiểu này, nếu không S2 dò ra 0 field và không điền được gì.
#
# KHÔNG nhận `___` (gạch dưới dài): trong mẫu đó là **dòng ký tên**, không phải chỗ điền.
#
# Bắt cả dãy TRỘN `…….. ` (vừa ký tự … vừa dấu chấm) làm MỘT chỗ trống. Nếu chỉ bắt riêng
# từng loại thì "Thời gian giao hàng: …….. giờ" điền xong còn sót ".." lơ lửng giữa câu.
BLANK = re.compile(r"[…\.]{2,}")


def _is_blank(tok):
    """Dãy dấu này là CHỖ TRỐNG hay chỉ là dấu câu?

    "…" hoặc "…." → chỗ trống. "..." (3 dấu chấm thường) → dấu lược, để nguyên. Cùng một
    phép kiểm dùng cho cả hàm dò và hàm điền, nếu lệch nhau thì giá trị rơi sai chỗ.
    """
    return "…" in tok or len(tok) >= 4
CANCEL_WORDS = ("huỷ", "huy", "thôi", "thoi", "dừng", "dung", "cancel")


# ---------------- đọc/điền docx (không cần python-docx cho phần dò) ----------------

def placeholders_in_docx(data):
    """Dò placeholder bằng cách đọc XML trong file docx — không phụ thuộc python-docx.

    Word hay cắt một chuỗi thành nhiều <w:t>, nên phải gộp text của cả document.xml lại
    rồi mới regex; nếu không sẽ bỏ sót đúng những placeholder bị cắt.
    """
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            xml = z.read("word/document.xml").decode("utf-8", "replace")
    except (zipfile.BadZipFile, KeyError):
        return []
    text = re.sub(r"<[^>]+>", "", xml)
    seen, out = set(), []
    for k in PLACEHOLDER.findall(text):
        if k not in seen:
            seen.add(k)
            out.append(k)
    return out


def _iter_paragraphs(doc):
    """Đi qua mọi paragraph theo MỘT thứ tự cố định, kèm ngữ cảnh → `(paragraph, ctx)`.

    Không phải thứ tự đọc thật (python-docx không cho biết bảng nằm giữa đoạn nào), nhưng
    **ổn định** — và ổn định là điều kiện duy nhất cần: hàm dò và hàm điền phải đi cùng
    một đường, nếu lệch thì giá trị rơi vào chỗ trống khác. Với hợp đồng thì đó là lỗi
    nguy hiểm nhất có thể có.

    Hai chi tiết của bảng, cả hai đều đã gây lỗi thật khi thử mẫu Mua bán:

    1. **Ô GỘP bị python-docx trả về NHIỀU LẦN** (`row.cells` lặp lại cùng một ô cho mỗi
       cột nó trải qua). Không lọc thì một chỗ trống bị đếm thành ba ⇒ lệch số thứ tự và
       **mọi giá trị phía sau rơi sai ô**.
    2. Bảng thông tin bên A/B theo layout `nhãn | : | giá trị`, nên ngữ cảnh đúng của một
       ô là **ô đầu hàng** ("Mã số thuế", "Địa chỉ trụ sở") — không có nó thì nhãn chỉ là
       "chỗ trống 11", người điền không biết điền gì.
    """
    for p in doc.paragraphs:
        yield p, ""
    for t in doc.tables:
        for row in t.rows:
            seen = set()
            head = ""
            for cell in row.cells:
                if id(cell._tc) in seen:
                    continue                       # ô gộp — đã đi qua rồi
                seen.add(id(cell._tc))
                if not head:
                    head = re.sub(r"\s+", " ", cell.text).strip()[:60]
                for p in cell.paragraphs:
                    yield p, head


def _blank_label(text, start, end, n):
    """Nhãn cho một chỗ trống = mảnh câu quanh nó, để người điền biết đang điền cái gì.

    Cắt theo BIÊN TỪ: cắt cứng 60 ký tự cho ra nhãn kiểu "ợp Đồng Mua Bán Hàng hóa…",
    đọc lên tưởng lỗi font. Thà nhãn ngắn hơn mà đọc được.
    """
    cut = max(0, start - 60)
    before = re.sub(r"\s+", " ", text[cut:start]).strip()
    if cut > 0:
        before = re.sub(r"^\S+\s+", "", before, count=1) or before
    after = re.sub(r"\s+", " ", text[end:end + 25]).strip()
    label = f"{before} ___ {after}".strip(" _")
    return label[:110] or f"chỗ trống {n}"


def blanks_in_docx(data):
    """Dò chỗ trống dạng `………` → list field theo thứ tự tài liệu.

    Trả `[{"key": "cho_trong_1", "label": "<mảnh câu quanh chỗ trống>", ...}]`. Nhãn quan
    trọng hơn key: "Thời gian giao hàng: ___ giờ ngày" thì người điền hiểu ngay, còn
    "cho_trong_7" thì không.
    """
    from docx import Document
    out = []
    try:
        doc = Document(io.BytesIO(data))
    except Exception:
        return []
    for p, ctx in _iter_paragraphs(doc):
        text = p.text
        if not text:
            continue
        for m in BLANK.finditer(text):
            if not _is_blank(m.group(0)):
                continue
            n = len(out) + 1
            lbl = _blank_label(text, m.start(), m.end(), n)
            if ctx and ctx not in lbl:
                lbl = f"{ctx} — {lbl}" if lbl != f"chỗ trống {n}" else ctx
            out.append({"key": f"cho_trong_{n}", "required": False, "label": lbl[:130]})
    return out


def fill_docx(data, values, mark=DRAFT_MARK):
    """Thay placeholder trong docx và đóng dấu DRAFT. Trả bytes file mới.

    Dùng python-docx để giữ định dạng. Word cắt placeholder qua nhiều `run` nên phải
    ghép text ở mức paragraph rồi mới thay — thay từng run sẽ trượt.
    """
    from docx import Document

    doc = Document(io.BytesIO(data))

    def sub(s):
        return PLACEHOLDER.sub(lambda m: str(values.get(m.group(1), m.group(0))), s)

    def fix_paragraph(p):
        if "{{" not in p.text:
            return
        new = sub(p.text)
        if new == p.text:
            return
        for r in p.runs[1:]:
            r.text = ""
        if p.runs:
            p.runs[0].text = new
        else:
            p.add_run(new)

    # Chỗ trống dạng `………`: đánh số theo ĐÚNG thứ tự `blanks_in_docx()` đã dò.
    # Field nào không có giá trị thì **để nguyên dấu ba chấm** — cùng luật với `{{...}}`:
    # xoá âm thầm là tạo ra hợp đồng thiếu điều khoản mà trông như đã hoàn chỉnh.
    counter = [0]

    def fix_blanks(p):
        if not BLANK.search(p.text):
            return
        def one(m):
            if not _is_blank(m.group(0)):
                return m.group(0)          # dấu lược, không phải chỗ trống
            counter[0] += 1
            v = values.get(f"cho_trong_{counter[0]}")
            return str(v) if v not in (None, "") else m.group(0)
        new = BLANK.sub(one, p.text)
        if new == p.text:
            return
        for r in p.runs[1:]:
            r.text = ""
        if p.runs:
            p.runs[0].text = new
        else:
            p.add_run(new)

    for p, _ctx in _iter_paragraphs(doc):
        fix_paragraph(p)
        fix_blanks(p)
    if mark:
        doc.add_paragraph("")
        doc.add_paragraph(f"[{mark}]")
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ---------------- registry ----------------

# Legal team đánh dấu bản MẪU bằng `[MAU]` / `[MẪU]` trong tên file. Dùng đúng quy ước
# của họ thay vì tự đoán: trong cùng folder còn có báo giá, đề nghị thanh toán, biên bản
# nghiệm thu thanh lý — toàn bộ đều có chữ "hợp đồng" trong tên. Lấy bừa theo từ khoá
# "hợp đồng" là agent đem *biên bản nghiệm thu* ra soạn thành hợp đồng.
TEMPLATE_MARK = ("[mau]", "[mẫu]")
# Folder KHÔNG phải kho mẫu, dù nằm trong folder mẫu.
SKIP_FOLDERS = ("ban thao", "bản thảo", "draft", "header", "footer")


def _is_template(name):
    low = (name or "").lower()
    return low.endswith(".docx") and any(m in low for m in TEMPLATE_MARK)


def _kind_of(folder_name):
    """"BỘ MẪU HỢP ĐỒNG_Mua bán" → "Mua bán". Dùng làm nhãn loại hợp đồng."""
    n = (folder_name or "").strip()
    return n.split("_", 1)[1].strip() if "_" in n else n


def _collect(lark, tokens, log):
    """File mẫu trong các folder đã khai **và một tầng folder con**.

    Vì sao phải xuống một tầng: legal team xếp mẫu theo loại
    (`BỘ MẪU HỢP ĐỒNG_Mua bán`, `_Dịch vụ`, `_Thuê nhà`…), nên ở tầng gốc không có file
    nào. Chỉ **một** tầng, không đệ quy sâu — và bỏ hẳn các folder ở `SKIP_FOLDERS`,
    quan trọng nhất là folder bản thảo agent tự xuất: quét vào đó là bản thảo của chính
    mình thành mẫu cho lần sau.
    """
    out, seen, skipped = [], set(), []

    def scan(tok, label, depth):
        try:
            entries = lark.drive_files(tok)
        except Exception as exc:
            log(f"[template] không đọc được folder {label or tok[:12]}: {exc}")
            return
        for f in entries:
            name = f.get("name") or ""
            if f.get("type") == "folder":
                if any(k in name.lower() for k in SKIP_FOLDERS):
                    continue
                if depth == 0:
                    scan(f["token"], _kind_of(name), 1)
                continue
            if f.get("token") in seen:
                continue
            seen.add(f["token"])
            if _is_template(name):
                out.append((f, label))
            elif name.lower().endswith((".docx", ".doc")):
                skipped.append(f"{label}/{name}" if label else name)
    for tok in tokens:
        scan(tok, "", 0)
    if skipped:
        log(f"[template] bỏ qua {len(skipped)} file không có dấu [MAU]: "
            + "; ".join(skipped[:4]) + ("…" if len(skipped) > 4 else ""))
    return out


def sync_templates(lark, store, folder_token, log=print):
    """Nạp registry template từ Drive → bảng contract_templates.

    `folder_token` nhận **nhiều folder**, cách nhau bằng dấu phẩy: trong Drive pháp chế có
    hai folder tên đều hợp lý cho mẫu hợp đồng, đoán sai một cái là S2 không thấy mẫu nào
    mà vẫn báo thành công.
    """
    tokens = [t.strip() for t in (folder_token or "").split(",") if t.strip()]
    if not tokens:
        return {"templates": 0, "skipped": "chưa cấu hình LEGAL_TEMPLATE_FOLDER"}
    found = _collect(lark, tokens, log)
    specs = {}      # <tên>.fields.json nằm cạnh file mẫu
    n = 0
    for f, kind in found:
        name = f.get("name", "")
        tok, stem = f.get("token"), name.rsplit(".", 1)[0]
        fields = None
        spec = specs.get(stem + ".fields.json")
        if spec:
            try:
                fields = json.loads(lark.drive_download(spec["token"]).decode("utf-8"))
            except Exception as exc:
                log(f"[template] {name}: .fields.json lỗi ({exc}) → dò placeholder")
        if fields is None:
            try:
                raw = lark.drive_download(tok)
            except Exception as exc:
                log(f"[template] {name}: không tải được ({exc}) → bỏ qua")
                continue
            keys = placeholders_in_docx(raw)
            if keys:
                fields = [{"key": k, "label": k.replace("_", " "), "required": True}
                          for k in keys]
            else:
                # Mẫu thật của legal team không có `{{...}}` — chỗ cần điền là dãy `………`
                # giữa câu. Không rơi về nhánh này thì mọi mẫu đều 0 field và S2 không
                # điền được gì, mà vẫn báo "đã nạp 9 mẫu".
                fields = blanks_in_docx(raw)
        label = f"{kind} · {stem}" if kind else stem
        store.write(
            "INSERT INTO contract_templates (key, name, file_token, lark_url, fields, "
            "edit_ts, status, updated_at) VALUES (?,?,?,?,?,?,'active',?) "
            "ON CONFLICT(key) DO UPDATE SET name=excluded.name, fields=excluded.fields, "
            "edit_ts=excluded.edit_ts, status='active', updated_at=excluded.updated_at",
            (f"drive:{tok}", label, tok, f.get("url") or lark.drive_file_url(tok),
             json.dumps(fields, ensure_ascii=False), f.get("modified_time"), time.time()))
        n += 1
        log(f"[template] {label}: {len(fields)} field")
    if not n:
        log(f"[template] {len(tokens)} folder, 0 file mẫu có dấu [MAU] — legal team chưa "
            f"bỏ mẫu vào, hoặc mẫu không đánh dấu [MAU] trong tên file")
    return {"templates": n}


def templates(store):
    rows = store.query("SELECT * FROM contract_templates WHERE status='active' ORDER BY name")
    for r in rows:
        try:
            r["fields"] = json.loads(r.get("fields") or "[]")
        except json.JSONDecodeError:
            r["fields"] = []
    return rows


# Từ chung có trong hầu hết tên mẫu → không mang thông tin phân biệt. Nếu tính cả
# những từ này thì mọi câu chứa "hợp đồng" đều khớp mẫu đầu tiên trong kho.
GENERIC_WORDS = {"hop", "dong", "hd", "mau", "ban", "va", "voi", "cua"}


def pick_template(store, text):
    """Khớp mẫu theo phần ĐẶC TRƯNG của tên (bỏ 'hợp đồng', không phân biệt dấu).

    Chỉ trả về khi có đúng MỘT mẫu khớp trọn phần đặc trưng. Mơ hồ → None để Agent liệt
    kê cho người chọn, thay vì đoán sai loại hợp đồng.
    """
    from legalkb.gates import plain
    want = plain(text)
    hits = []
    for t in templates(store):
        words = [w for w in plain(t["name"]).split()
                 if w not in GENERIC_WORDS and len(w) > 1]
        if words and all(w in want for w in words):
            hits.append(t)
    return hits[0] if len(hits) == 1 else None


# ---------------- state đa lượt ----------------

class Drafts:
    def __init__(self, store):
        self.store = store

    def get(self, session_id):
        d = self.store.one("SELECT * FROM contract_drafts WHERE session_id=?", (session_id,))
        if d:
            try:
                d["values"] = json.loads(d.get("values_json") or "{}")
            except json.JSONDecodeError:
                d["values"] = {}
        return d

    def save(self, session_id, **f):
        cur = self.get(session_id)
        if "values" in f:
            f["values_json"] = json.dumps(f.pop("values"), ensure_ascii=False)
        f["updated_at"] = time.time()
        if cur:
            sets = ", ".join(f"{k}=?" for k in f)
            self.store.write(f"UPDATE contract_drafts SET {sets} WHERE session_id=?",
                             (*f.values(), session_id))
        else:
            cols = ["session_id", *f.keys()]
            self.store.write(
                f"INSERT INTO contract_drafts ({', '.join(cols)}) "
                f"VALUES ({', '.join('?' * len(cols))})", (session_id, *f.values()))
        return self.get(session_id)

    def drop(self, session_id):
        self.store.write("DELETE FROM contract_drafts WHERE session_id=?", (session_id,))


def missing_fields(template, values):
    return [f for f in template["fields"]
            if f.get("required", True) and not str(values.get(f["key"], "")).strip()]


def ask_next(template, values):
    """Câu hỏi cho field còn thiếu tiếp theo, hoặc None nếu đã đủ."""
    miss = missing_fields(template, values)
    if not miss:
        return None, None
    f = miss[0]
    q = f"Cho mình **{f.get('label') or f['key']}**?"
    if f.get("hint"):
        q += f"\n_({f['hint']})_"
    left = len(miss) - 1
    if left:
        q += f"\n\n_Còn {left} thông tin nữa._"
    return f["key"], q


def summary(template, values):
    lines = [f"**{template['name']}** — xác nhận thông tin:"]
    for f in template["fields"]:
        v = values.get(f["key"], "")
        lines.append(f"- {f.get('label') or f['key']}: **{v or '(trống)'}**")
    lines.append("\nĐúng thì trả lời **ok**; muốn sửa thì ghi ví dụ "
                 "`sửa gia_tri: 500 triệu`; bỏ thì ghi **huỷ**.")
    return "\n".join(lines)


EDIT_RE = re.compile(r"^\s*(?:sửa|sua)\s+([a-zA-Z0-9_]+)\s*[:=]\s*(.+)$", re.I | re.S)


def parse_edit(text):
    m = EDIT_RE.match(text or "")
    return (m.group(1), m.group(2).strip()) if m else (None, None)


def wants_cancel(text):
    t = (text or "").strip().lower()
    return any(t == w or t.startswith(w + " ") for w in CANCEL_WORDS)


_FEEDBACK_PROMPT = """Góp ý của bộ phận Pháp chế về bản thảo hợp đồng cần được quy về việc
sửa giá trị của các field. Trả về DUY NHẤT một JSON:

{"updates": {"<field_key>": "<giá trị mới>"}, "unresolved": "phần góp ý KHÔNG quy được về
field nào (chuỗi rỗng nếu không có)"}

Chỉ dùng đúng các field_key có trong danh sách. Không chắc thì để vào unresolved, TUYỆT
ĐỐI không đoán giá trị.

Field hiện có: {fields}
Giá trị hiện tại: {values}
Góp ý của Pháp chế: {comment}
"""


def apply_feedback(brain, template, values, comment, model=None):
    """Quy góp ý của Pháp chế về việc sửa field.

    Trả (values_mới, unresolved). Không quy được thì `unresolved` giữ nguyên văn góp ý —
    để Agent hỏi lại người yêu cầu thay vì tự đoán giá trị hợp đồng.
    """
    import json as _json
    import re as _re
    fields = ", ".join(f"{f['key']} ({f.get('label') or f['key']})"
                       for f in template["fields"])
    raw = brain.call_claude(
        _FEEDBACK_PROMPT.replace("{fields}", fields)
                        .replace("{values}", _json.dumps(values, ensure_ascii=False))
                        .replace("{comment}", comment or ""),
        model=model, timeout=90)
    m = _re.search(r"\{.*\}", raw or "", _re.S)
    if not m:
        return dict(values), (comment or "")
    try:
        out = _json.loads(m.group(0))
    except _json.JSONDecodeError:
        return dict(values), (comment or "")
    valid = {f["key"] for f in template["fields"]}
    new = dict(values)
    for k, v in (out.get("updates") or {}).items():
        if k in valid and str(v).strip():
            new[k] = str(v).strip()
    return new, (out.get("unresolved") or "").strip()
