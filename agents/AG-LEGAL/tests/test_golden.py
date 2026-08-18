"""Test phép kiểm chống bịa nguồn + trích text file (PLAN Phase 7 mục 44)."""
import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
os.environ.setdefault("NLM_NOTEBOOK_KB_ID", "nb-test")

import golden_run
from legalkb import extract
from tests.test_flows import docx_with

KNOWN = {"https://tenant/wiki/a", "https://tenant/file/b"}


def test_fabricated_citation_is_caught():
    """Link không có trong legal_sources = BỊA NGUỒN. Đây là phép kiểm quan trọng nhất."""
    answer = "Theo quy định...\n📎 **Nguồn:**\n- [Quy chế](https://tenant/wiki/KHONG-CO)"
    ok, errs = golden_run.grade(answer, {"must_cite": True}, KNOWN)
    assert not ok and any("BỊA NGUỒN" in e for e in errs)


def test_real_citation_passes():
    answer = "Theo quy định...\n📎 **Nguồn:**\n- [Quy chế](https://tenant/wiki/a)"
    ok, errs = golden_run.grade(answer, {"must_cite": True}, KNOWN)
    assert ok, errs


def test_missing_citation_fails_when_required():
    ok, errs = golden_run.grade("Theo quy định là 45 ngày.", {"must_cite": True}, KNOWN)
    assert not ok and "không có trích dẫn nào" in errs


def test_refusal_case_does_not_need_citation():
    ok, _ = golden_run.grade("Tài liệu hiện chưa quy định nội dung này.",
                             {"must_cite": False, "must_have": ["chưa"]}, KNOWN)
    assert ok


def test_must_not_have_catches_invented_authority():
    """Câu 'chưa có quy định' mà lại viện dẫn điều khoản là dấu hiệu bịa."""
    ok, errs = golden_run.grade("Theo quy định tại Điều 5 Nghị định 99...",
                                {"must_cite": False, "must_not_have": ["theo quy định tại"]},
                                KNOWN)
    assert not ok and any("không được có" in e for e in errs)


def test_empty_known_urls_skips_fabrication_check():
    """KB chưa sync thì không có gì đối chiếu — không được báo sai là bịa."""
    answer = "x\n📎 **Nguồn:**\n- [A](https://tenant/wiki/z)"
    ok, _ = golden_run.grade(answer, {"must_cite": True}, set())
    assert ok


def test_selfcheck_of_shipped_cases_passes():
    assert golden_run.selfcheck() == 0


# ---------------- trích text ----------------

def test_extract_docx_includes_table_cells():
    """Điều khoản hợp đồng hay nằm trong bảng — bỏ bảng là bỏ mất điều khoản."""
    from docx import Document
    import io
    d = Document()
    d.add_paragraph("Điều 1")
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Giá trị"
    t.rows[0].cells[1].text = "500 triệu"
    out = io.BytesIO()
    d.save(out)
    text = extract.from_bytes(out.getvalue(), "hd.docx")
    assert "Điều 1" in text and "500 triệu" in text


def test_extract_rejects_unknown_extension():
    try:
        extract.from_bytes(b"abc", "anh.png")
    except extract.ExtractError as e:
        assert "chưa đọc được định dạng" in str(e)
    else:
        raise AssertionError("phải báo lỗi định dạng")


def test_extract_rejects_empty_document():
    try:
        extract.from_bytes(docx_with(""), "trong.docx")
    except extract.ExtractError as e:
        assert "không có nội dung" in str(e)
    else:
        raise AssertionError("file rỗng phải báo lỗi, không trả text rỗng")
