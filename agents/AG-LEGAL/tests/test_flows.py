"""Test S2–S5 + hệ quả sau khi Pháp chế quyết định (PLAN Phase 3–6) — offline."""
import io
import json
import pathlib
import os
import sys

import pytest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
os.environ.setdefault("NLM_NOTEBOOK_KB_ID", "nb-test")

import consumer
from legalkb import contracts, flows, news, review as review_mod, signing
from legalkb.engine import EngineAnswer
from legalkb.gates import GATE
from tests.test_consumer import make

GROUP = consumer.GROUP_CHAT_ID


class FakeLarkKB:
    """Drive giả: download trả docx thật (dựng bằng python-docx), upload ghi lại."""

    def __init__(self, template_bytes=b"", files=None):
        self.template_bytes = template_bytes
        self.uploaded = []
        self._files = files or []

    def drive_files(self, folder):
        return self._files

    def drive_download(self, token):
        return self.template_bytes

    def drive_upload(self, folder, name, data):
        self.uploaded.append((folder, name, data))
        return "tok_" + str(len(self.uploaded))

    def drive_file_url(self, token, file_type="file"):
        return f"https://tenant/file/{token}"


def docx_with(text):
    from docx import Document
    d = Document()
    d.add_paragraph(text)
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def add_template(store, name="Hợp đồng dịch vụ", fields=None):
    fields = fields or [{"key": "ten_ben_a", "label": "tên bên A", "required": True},
                        {"key": "gia_tri", "label": "giá trị hợp đồng", "required": True}]
    store.write(
        "INSERT INTO contract_templates (key, name, file_token, lark_url, fields, status) "
        "VALUES ('drive:t1',?, 't1','https://tenant/file/t1',?,'active')",
        (name, json.dumps(fields, ensure_ascii=False)))


def add_approver(store):
    store.write("INSERT INTO legal_roles (email, role, contract_type, open_id, name, active)"
                " VALUES ('thint@hapas.vn','approver','*','ou_thint','Thi',1)")


# ==================== S2: điền docx ====================

def test_placeholders_found_even_when_word_splits_them():
    """Word cắt một chuỗi thành nhiều <w:t>; dò theo từng run sẽ bỏ sót placeholder."""
    data = docx_with("Bên A: {{ten_ben_a}}, giá trị {{gia_tri}} đồng")
    assert set(contracts.placeholders_in_docx(data)) == {"ten_ben_a", "gia_tri"}


def test_fill_docx_replaces_and_stamps_draft():
    from docx import Document
    data = docx_with("Bên A: {{ten_ben_a}} — giá trị {{gia_tri}}")
    out = contracts.fill_docx(data, {"ten_ben_a": "Công ty ABC", "gia_tri": "500 triệu"})
    text = "\n".join(p.text for p in Document(io.BytesIO(out)).paragraphs)
    assert "Công ty ABC" in text and "500 triệu" in text
    assert "{{" not in text
    assert contracts.DRAFT_MARK in text          # bản thảo luôn có dấu DRAFT


def test_fill_docx_keeps_unknown_placeholder_visible():
    """Field thiếu thì để nguyên {{...}} — im lặng xoá đi là cách tạo hợp đồng thiếu điều khoản."""
    out = contracts.fill_docx(docx_with("X: {{chua_co}}"), {})
    from docx import Document
    assert "{{chua_co}}" in "\n".join(p.text for p in Document(io.BytesIO(out)).paragraphs)


# ==================== S2: luồng đa lượt ====================

def test_s2_asks_fields_one_by_one_then_confirms(tmp_path):
    store, pf, eng, g, b = make(tmp_path)
    add_template(store)
    out1 = flows.s2_create(b, "tạo hợp đồng dịch vụ", "s1", "ou_u", "oc_u")
    assert "tên bên A" in out1
    out2 = flows.s2_create(b, "Công ty ABC", "s1", "ou_u", "oc_u")
    assert "giá trị hợp đồng" in out2 and "Còn" not in out2.split("\n")[0]
    out3 = flows.s2_create(b, "500 triệu", "s1", "ou_u", "oc_u")
    assert "xác nhận thông tin" in out3 and "Công ty ABC" in out3


def test_s2_state_survives_restart(tmp_path):
    """Bỏ dở rồi quay lại (Bundle mới = tiến trình mới) vẫn hỏi tiếp đúng field."""
    store, pf, eng, g, b = make(tmp_path)
    add_template(store)
    flows.s2_create(b, "tạo hợp đồng dịch vụ", "s-keep", "ou_u", "oc_u")
    flows.s2_create(b, "Công ty ABC", "s-keep", "ou_u", "oc_u")
    b2 = flows.Bundle(pf=pf, store=store, engine=eng, gates=g)   # "restart"
    assert "giá trị hợp đồng" in flows.s2_create(b2, "", "s-keep", "ou_u", "oc_u")


def test_s2_edit_before_confirm(tmp_path):
    store, pf, eng, g, b = make(tmp_path)
    add_template(store)
    for msg in ("tạo hợp đồng dịch vụ", "Công ty ABC", "500 triệu"):
        flows.s2_create(b, msg, "s2", "ou_u", "oc_u")
    out = flows.s2_create(b, "sửa gia_tri: 800 triệu", "s2", "ou_u", "oc_u")
    assert "800 triệu" in out


def test_s2_cancel(tmp_path):
    store, pf, eng, g, b = make(tmp_path)
    add_template(store)
    flows.s2_create(b, "tạo hợp đồng dịch vụ", "s3", "ou_u", "oc_u")
    assert "Đã bỏ" in flows.s2_create(b, "huỷ", "s3", "ou_u", "oc_u")
    assert b.drafts.get("s3") is None


def test_s2_unknown_template_lists_options_without_inventing(tmp_path):
    store, pf, eng, g, b = make(tmp_path)
    add_template(store)
    out = flows.s2_create(b, "tạo hợp đồng thuê máy bay", "s4", "ou_u", "oc_u")
    assert "Hợp đồng dịch vụ" in out
    assert "máy bay" not in out            # không bịa ra mẫu không có


def test_s2_draft_goes_to_gate_not_to_requester(tmp_path, monkeypatch):
    """Review §A.1: bản thảo KHÔNG được gửi thẳng người yêu cầu."""
    lark = FakeLarkKB(docx_with("Bên A: {{ten_ben_a}} — {{gia_tri}}"))
    store, pf, eng, g, b = make(tmp_path, lark=lark)
    add_template(store)
    monkeypatch.setenv(flows.DRAFT_FOLDER_ENV, "fld_draft")
    for msg in ("tạo hợp đồng dịch vụ", "Công ty ABC", "500 triệu", "ok"):
        out = flows.s2_create(b, msg, "s5", "ou_u", "oc_requester")
    assert "chuyển bộ phận Pháp chế kiểm tra" in out
    assert lark.uploaded and lark.uploaded[0][1].startswith("DRAFT-")
    gate = store.one("SELECT * FROM legal_gates WHERE kind='s2_draft'")
    assert gate and gate["level"] == GATE and gate["status"] == "open"
    # người yêu cầu CHƯA nhận link nào
    assert not any(to == "oc_requester" for to, _ in pf.sent)


def test_s2_without_draft_folder_says_so(tmp_path, monkeypatch):
    lark = FakeLarkKB(docx_with("{{ten_ben_a}}"))
    store, pf, eng, g, b = make(tmp_path, lark=lark)
    add_template(store, fields=[{"key": "ten_ben_a", "label": "tên bên A"}])
    monkeypatch.delenv(flows.DRAFT_FOLDER_ENV, raising=False)
    flows.s2_create(b, "tạo hợp đồng dịch vụ", "s6", "ou_u", "oc_u")
    flows.s2_create(b, "ABC", "s6", "ou_u", "oc_u")
    out = flows.s2_create(b, "ok", "s6", "ou_u", "oc_u")
    assert "chưa xuất được file" in out and flows.DRAFT_FOLDER_ENV in out


# ==================== S2: sau khi Pháp chế quyết ====================

def _draft_to_gate(tmp_path, monkeypatch):
    lark = FakeLarkKB(docx_with("Bên A: {{ten_ben_a}} — {{gia_tri}}"))
    store, pf, eng, g, b = make(tmp_path, lark=lark)
    add_template(store)
    add_approver(store)
    monkeypatch.setenv(flows.DRAFT_FOLDER_ENV, "fld_draft")
    for msg in ("tạo hợp đồng dịch vụ", "Công ty ABC", "500 triệu", "ok"):
        flows.s2_create(b, msg, "sess", "ou_u", "oc_requester")
    return store, pf, g, b, store.one("SELECT * FROM legal_gates WHERE kind='s2_draft'")


def test_approved_draft_is_sent_to_requester(tmp_path, monkeypatch):
    store, pf, g, b, gate = _draft_to_gate(tmp_path, monkeypatch)
    pf.sent.clear()
    out = consumer.handle_group(
        {"id": 1, "payload": {"text": f"#{gate['id']} duyệt", "chat_id": GROUP,
                              "sender_open_id": "ou_thint"}}, b)
    assert "DUYỆT" in out
    msg = next(m for to, m in pf.sent if to == "oc_requester")
    assert "qua Pháp chế kiểm tra" in msg and "https://tenant/file/" in msg
    assert b.drafts.get("sess")["status"] == "done"


def test_rejected_draft_tells_requester_the_reason(tmp_path, monkeypatch):
    store, pf, g, b, gate = _draft_to_gate(tmp_path, monkeypatch)
    pf.sent.clear()
    consumer.handle_group(
        {"id": 2, "payload": {"text": f"#{gate['id']} huỷ: sai pháp nhân",
                              "chat_id": GROUP, "sender_open_id": "ou_thint"}}, b)
    msg = next(m for to, m in pf.sent if to == "oc_requester")
    assert "không thông qua" in msg and "sai pháp nhân" in msg
    assert b.drafts.get("sess")["status"] == "cancelled"


def test_changes_request_applies_feedback_and_reopens_gate(tmp_path, monkeypatch):
    store, pf, g, b, gate = _draft_to_gate(tmp_path, monkeypatch)
    monkeypatch.setattr(flows.brain, "call_claude", lambda *a, **k:
                        '{"updates": {"gia_tri": "800 triệu"}, "unresolved": ""}')
    consumer.handle_group(
        {"id": 3, "payload": {"text": f"#{gate['id']} sửa: giá trị phải là 800 triệu",
                              "chat_id": GROUP, "sender_open_id": "ou_thint"}}, b)
    assert b.drafts.get("sess")["values"]["gia_tri"] == "800 triệu"
    gates = store.query("SELECT * FROM legal_gates WHERE kind='s2_draft' ORDER BY id")
    assert len(gates) == 2 and gates[1]["round"] == 2      # gate mới, vòng 2


def test_feedback_that_cannot_map_asks_human_instead_of_guessing(tmp_path, monkeypatch):
    store, pf, g, b, gate = _draft_to_gate(tmp_path, monkeypatch)
    monkeypatch.setattr(flows.brain, "call_claude", lambda *a, **k:
                        '{"updates": {}, "unresolved": "điều khoản bảo mật chưa ổn"}')
    pf.sent.clear()
    out = consumer.handle_group(
        {"id": 4, "payload": {"text": f"#{gate['id']} sửa: điều khoản bảo mật chưa ổn",
                              "chat_id": GROUP, "sender_open_id": "ou_thint"}}, b)
    assert "không đoán" in out
    assert any("sửa <tên_field>" in m for to, m in pf.sent if to == "oc_requester")
    assert len(store.query("SELECT * FROM legal_gates WHERE kind='s2_draft'")) == 1


# ==================== S3: review hợp đồng ====================

def test_s3_asks_for_file_when_none_attached(tmp_path):
    store, pf, eng, g, b = make(tmp_path)
    out = flows.s3_review(b, {"id": 1, "payload": {}}, "nhờ xem hợp đồng", "s", "ou", "oc")
    assert "đính kèm file" in out


def test_s3_reports_findings_and_does_not_close_when_dirty(tmp_path, monkeypatch):
    store, pf, eng, g, b = make(tmp_path,
                                engine_answer=EngineAnswer(ok=True, text="Checklist: A, B"))
    pf.resource = docx_with("Điều 1. Giá trị hợp đồng")
    monkeypatch.setattr(pf, "lark_resource", lambda *a, **k: pf.resource, raising=False)
    monkeypatch.setattr(flows.brain, "call_claude", lambda *a, **k: json.dumps(
        {"findings": [{"severity": "high", "clause": "THIẾU", "issue": "thiếu điều khoản phạt",
                       "suggestion": "bổ sung mức phạt"}], "clean": False, "note": "cần sửa"}))
    job = {"id": 2, "payload": {"file_key": "fk", "file_name": "hd.docx",
                                "message_id": "om_1", "chat_id": "oc_u"}}
    out = flows.s3_review(b, job, "nhờ review", "s3a", "ou_u", "oc_u")
    assert "thiếu điều khoản phạt" in out and "gửi lại file" in out
    r = b.reviews.latest("s3a")
    assert r["status"] == "issues_sent"
    assert not store.query("SELECT * FROM legal_gates WHERE kind='s3_review' AND level='gate'")


def test_s3_clean_contract_goes_to_approver(tmp_path, monkeypatch):
    store, pf, eng, g, b = make(tmp_path,
                                engine_answer=EngineAnswer(ok=True, text="Checklist: A"))
    monkeypatch.setattr(pf, "lark_resource", lambda *a, **k: docx_with("Điều 1"), raising=False)
    monkeypatch.setattr(flows.brain, "call_claude", lambda *a, **k: json.dumps(
        {"findings": [], "clean": True, "note": "ổn"}))
    job = {"id": 3, "payload": {"file_key": "fk", "file_name": "hd.docx",
                                "message_id": "om", "chat_id": "oc_u"}}
    out = flows.s3_review(b, job, "review", "s3b", "ou_u", "oc_u")
    assert "chuyển bộ phận Pháp chế xác nhận" in out
    assert b.reviews.latest("s3b")["status"] == "pending_approval"
    assert store.one("SELECT * FROM legal_gates WHERE kind='s3_review' AND level='gate'")


def test_s3_model_failure_never_declares_contract_clean(tmp_path, monkeypatch):
    """Model lỗi thì KHÔNG được kết luận 'hợp đồng sạch' — điểm an toàn."""
    monkeypatch.setattr(flows.brain, "call_claude", lambda *a, **k: "xin lỗi tôi không thể")
    findings, clean, note = review_mod.analyse(flows.brain, "hợp đồng...", "checklist", "x.pdf")
    assert findings == [] and clean is False
    assert "chưa rà soát được" in note.lower()


def test_s3_report_warns_when_kb_has_no_checklist():
    out = review_mod.render_report([], True, "", "hd.pdf", has_checklist=False)
    assert "chưa có" in out and "nguyên tắc chung" in out


def test_s3_unreadable_file_says_so(tmp_path, monkeypatch):
    store, pf, eng, g, b = make(tmp_path)
    monkeypatch.setattr(pf, "lark_resource", lambda *a, **k: b"khong-phai-docx", raising=False)
    job = {"id": 4, "payload": {"file_key": "fk", "file_name": "anh.png",
                                "message_id": "om", "chat_id": "oc_u"}}
    out = flows.s3_review(b, job, "review", "s3c", "ou_u", "oc_u")
    assert "chưa đọc được file" in out


# ==================== S4: crawl + digest ====================

def test_rss_parsed_and_doc_number_extracted():
    xml = """<rss><channel>
      <item><title>Nghị định 15/2026/NĐ-CP về thương mại điện tử</title>
            <link>https://x.vn/a</link><description>Nội dung</description></item>
      <item><title>Thông tư mới</title><link>https://x.vn/b</link></item>
    </channel></rss>"""
    items = news.parse_rss(xml)
    assert len(items) == 2
    assert news.doc_no_of(items[0]["title"]) == "15/2026/NĐ-CP"
    assert news.doc_no_of(items[1]["title"]) is None


def test_crawl_dedupes_by_doc_number(tmp_path, monkeypatch):
    store, pf, eng, g, b = make(tmp_path)
    store.write("INSERT INTO legal_news_sources (name, url, kind, active) "
                "VALUES ('X','https://x.vn/rss','rss',1)")
    xml = ("<rss><channel><item><title>Nghị định 15/2026/NĐ-CP</title>"
           "<link>https://x.vn/a</link></item></channel></rss>")
    monkeypatch.setattr(news, "fetch", lambda *a, **k: xml)
    assert len(news.crawl(store, log=lambda m: None)) == 1
    assert news.crawl(store, log=lambda m: None) == []      # lần 2 không trùng


def test_one_broken_source_does_not_kill_others(tmp_path, monkeypatch):
    store, pf, eng, g, b = make(tmp_path)
    store.write("INSERT INTO legal_news_sources (name,url,kind,active) "
                "VALUES ('Vỡ','https://bad.vn/rss','rss',1)")
    store.write("INSERT INTO legal_news_sources (name,url,kind,active) "
                "VALUES ('Tốt','https://ok.vn/rss','rss',1)")

    def fetch(url, timeout=30):
        if "bad" in url:
            raise RuntimeError("403 chặn bot")
        return ("<rss><channel><item><title>Nghị định 9/2026/NĐ-CP</title>"
                "<link>https://ok.vn/a</link></item></channel></rss>")

    monkeypatch.setattr(news, "fetch", fetch)
    assert len(news.crawl(store, log=lambda m: None)) == 1
    bad = store.one("SELECT * FROM legal_news_sources WHERE name='Vỡ'")
    assert "403" in bad["last_error"]


def test_item_without_source_link_is_dropped(tmp_path, monkeypatch):
    """Review §B: mục không trích được link nguồn thì loại, không đưa vào digest."""
    store, pf, eng, g, b = make(tmp_path)
    store.write("INSERT INTO legal_news_items (key, url, title, status, found_at) "
                "VALUES ('k1','','Không có link','new',0)")
    monkeypatch.setattr(flows.brain, "call_claude", lambda *a, **k: "{}")
    assert news.summarise(store, flows.brain, ["k1"], log=lambda m: None) == []
    assert store.one("SELECT status FROM legal_news_items WHERE key='k1'")["status"] == "dropped"


def test_digest_opens_gate_and_publishes_nothing_yet(tmp_path, monkeypatch):
    """Review §B: chưa duyệt thì group không nhận digest, notebook không có source mới."""
    store, pf, eng, g, b = make(tmp_path)
    store.write("INSERT INTO legal_news_sources (name,url,kind,active) "
                "VALUES ('X','https://x.vn/rss','rss',1)")
    monkeypatch.setattr(news, "fetch", lambda *a, **k:
                        "<rss><channel><item><title>Nghị định 7/2026/NĐ-CP</title>"
                        "<link>https://x.vn/a</link></item></channel></rss>")
    monkeypatch.setattr(flows.brain, "call_claude", lambda *a, **k: json.dumps(
        {"scope": "TMĐT", "effective": "01/01/2027", "impact": "ảnh hưởng bán online"}))
    gid = flows.news_cycle(b, GROUP, log=lambda m: None)
    gate = b.gates.get(gid)
    assert gate["level"] == GATE and gate["status"] == "open"
    assert store.one("SELECT status FROM legal_news_items WHERE doc_no='7/2026/NĐ-CP'"
                     )["status"] == "in_digest"
    # card thông báo có gửi, nhưng digest thì CHƯA phát hành
    assert not any("Hiệu lực" in (m or "") for _, m in pf.sent)


def test_approved_digest_is_published_and_ingested(tmp_path, monkeypatch):
    store, pf, eng, g, b = make(tmp_path, lark=FakeLarkKB())
    add_approver(store)
    ingested = []
    eng.add_text_source = lambda title, content: ingested.append(title)
    store.write("INSERT INTO legal_news_sources (name,url,kind,active) "
                "VALUES ('X','https://x.vn/rss','rss',1)")
    monkeypatch.setattr(news, "fetch", lambda *a, **k:
                        "<rss><channel><item><title>Nghị định 8/2026/NĐ-CP</title>"
                        "<link>https://x.vn/a</link></item></channel></rss>")
    monkeypatch.setattr(flows.brain, "call_claude", lambda *a, **k: json.dumps(
        {"scope": "s", "effective": "e", "impact": "i"}))
    gid = flows.news_cycle(b, GROUP, log=lambda m: None)
    monkeypatch.setenv("LEGAL_DRIVE_FOLDER", "fld_law")
    pf.sent.clear()
    out = consumer.handle_group({"id": 9, "payload": {"text": f"#{gid} duyệt",
                                                      "chat_id": GROUP,
                                                      "sender_open_id": "ou_thint"}}, b)
    assert "phát hành digest" in out
    assert any("Hiệu lực" in (m or "") for _, m in pf.sent)      # digest đã gửi group
    assert ingested                                              # đã nạp notebook
    assert store.one("SELECT status FROM legal_news_items WHERE doc_no='8/2026/NĐ-CP'"
                     )["status"] == "published"


def test_rejected_digest_publishes_nothing(tmp_path, monkeypatch):
    store, pf, eng, g, b = make(tmp_path)
    add_approver(store)
    store.write("INSERT INTO legal_news_items (key, doc_no, url, title, summary, status, "
                "found_at) VALUES ('k9','9/2026/NĐ-CP','https://x/a','T','{}','in_digest',0)")
    gid = b.gates.open("s4_digest", GATE, payload={"keys": ["k9"], "digest": "d"})
    pf.sent.clear()
    out = consumer.handle_group({"id": 10, "payload": {"text": f"#{gid} huỷ: chưa cần",
                                                       "chat_id": GROUP,
                                                       "sender_open_id": "ou_thint"}}, b)
    assert "không gửi, không nạp KB" in out
    assert store.one("SELECT status FROM legal_news_items WHERE key='k9'")["status"] == "dropped"


# ==================== S5: trình ký ====================

def test_s5_step3_flags_missing_documents(tmp_path, monkeypatch):
    store, pf, eng, g, b = make(tmp_path)
    signing.set_checklist(store, "hợp đồng dịch vụ", ["Giấy ĐKKD", "Báo giá", "Tờ trình"])
    monkeypatch.setattr(flows.brain, "call_claude", lambda *a, **k: json.dumps(
        {"missing_docs": ["Báo giá"], "findings": [], "note": "thiếu 1 mục"}))
    res = signing.step3(flows.brain, store, "nội dung hồ sơ", "hợp đồng dịch vụ")
    out = signing.render_step3(res, "hợp đồng dịch vụ")
    assert res["ok"] and "Báo giá" in out and "Thiếu đầu mục" in out


def test_s5_step3_never_blocks_when_model_fails(tmp_path, monkeypatch):
    """Agent lỗi thì hồ sơ VẪN đi tiếp — máy không được làm nghẽn quy trình người."""
    store, pf, eng, g, b = make(tmp_path)
    monkeypatch.setattr(flows.brain, "call_claude", lambda *a, **k: "lỗi")
    res = signing.step3(flows.brain, store, "hồ sơ", "HĐ")
    assert res["ok"] is False
    assert "chưa rà soát kịp" in signing.render_step3(res)


def test_s5_step5_high_severity_bounces_back_to_step4(monkeypatch):
    monkeypatch.setattr(flows.brain, "call_claude", lambda *a, **k: json.dumps(
        {"findings": [{"severity": "high", "issue": "sai pháp nhân bên B"}], "note": ""}))
    res = signing.step5(flows.brain, "hồ sơ")
    assert res["blocking"] is True
    assert "quay lại Bước 4" in signing.render_step5(res, bounce_count=0)


def test_s5_step5_low_severity_is_advisory_only(monkeypatch):
    monkeypatch.setattr(flows.brain, "call_claude", lambda *a, **k: json.dumps(
        {"findings": [{"severity": "low", "issue": "câu chữ điều 3"}], "note": ""}))
    res = signing.step5(flows.brain, "hồ sơ")
    assert res["blocking"] is False
    out = signing.render_step5(res)
    assert "cảnh báo tham khảo" in out and "quay lại Bước 4" not in out


def test_s5_stops_bouncing_after_max_and_escalates(monkeypatch):
    monkeypatch.setattr(flows.brain, "call_claude", lambda *a, **k: json.dumps(
        {"findings": [{"severity": "high", "issue": "x"}], "note": ""}))
    res = signing.step5(flows.brain, "hồ sơ")
    out = signing.render_step5(res, bounce_count=signing.MAX_BOUNCE)
    assert "không quay lại" in out and "trưởng Pháp chế" in out


def test_s5_observe_gate_auto_passes_on_sla(tmp_path, monkeypatch):
    """Gate của S5 là observe → quá hạn tự thông, không giữ hồ sơ."""
    store, pf, eng, g, b = make(tmp_path)
    monkeypatch.setattr(pf, "lark_resource", lambda *a, **k: docx_with("hồ sơ"), raising=False)
    monkeypatch.setattr(flows.brain, "call_claude", lambda *a, **k: json.dumps(
        {"missing_docs": [], "findings": [], "note": "ok"}))
    job = {"id": 5, "payload": {"file_key": "fk", "file_name": "hs.docx",
                                "message_id": "om", "chat_id": "oc_u"}}
    flows.s5_dossier(b, job, "hồ sơ trình ký bước 3", "s5a", "ou_u", "oc_u")
    gate = store.one("SELECT * FROM legal_gates WHERE kind='s5_step3'")
    assert gate["level"] == "observe"
    store.write("UPDATE legal_gates SET sla_deadline=0 WHERE id=?", (gate["id"],))
    assert ("auto_passed", gate["id"]) in b.gates.sla_tick()


def test_s5_says_it_is_shadow_mode(tmp_path):
    store, pf, eng, g, b = make(tmp_path)
    out = flows.s5_dossier(b, {"id": 6, "payload": {}}, "trình ký", "s", "ou", "oc")
    assert "chạy thử song song" in out and "Lark Approval" in out


# ==================== S5: form thật của Lark Approval ====================

def test_parse_approval_form_maps_real_widget_ids():
    """Widget id đọc live từ workflow "Review và phê duyệt hợp đồng" (19/08/2026)."""
    form = json.dumps([
        {"id": "widget16500094298477484625790733479", "type": "input",
         "name": "Tên hợp đồng", "value": "HĐ dịch vụ ABC"},
        {"id": "widget16500094292359699894458073106", "type": "textarea",
         "name": "Mô tả nội dung cần review và phê duyệt", "value": "Nhờ rà soát điều 5"},
        {"id": "widget17871105849310001", "type": "date",
         "name": "Ngày cần hoàn thành", "value": "2026-09-01"},
    ])
    out = signing.parse_form(form)
    assert out["contract_name"] == "HĐ dịch vụ ABC"
    assert out["description"] == "Nhờ rà soát điều 5"
    assert out["due_date"] == "2026-09-01"
    assert out["attachments"] is None          # không đính kèm → None, không nổ
    assert len(out["raw"]) == 3                # giữ nguyên field lạ để không mất dữ liệu


def test_parse_approval_form_survives_garbage():
    assert signing.parse_form("không phải json") == {}
    assert signing.parse_form(None) == {"raw": [], "contract_name": None,
                                        "description": None, "attachments": None,
                                        "due_date": None}


# ==================== bộ nhớ index ====================

def test_approved_draft_is_indexed_into_brain(tmp_path, monkeypatch):
    """Hợp đồng đã gửi phải vào brain của platform để lượt sau tra lại được qua RAG."""
    store, pf, g, b, gate = _draft_to_gate(tmp_path, monkeypatch)
    items = []
    monkeypatch.setattr(pf, "add_brain_item",
                        lambda title, content, status="approved", source_url=None:
                        items.append((title, content, source_url)), raising=False)
    consumer.handle_group({"id": 30, "payload": {"text": f"#{gate['id']} duyệt",
                                                 "chat_id": GROUP,
                                                 "sender_open_id": "ou_thint"}}, b)
    assert items, "duyệt xong mà không index thì lượt sau không tra lại được"
    title, content, url = items[0]
    assert title.startswith("[Hợp đồng đã gửi]")
    assert "Người duyệt: thint@hapas.vn" in content
    assert url and url.startswith("https://")      # phải có link đối chứng


def test_template_content_indexed_not_just_name(tmp_path, monkeypatch):
    """Index NỘI DUNG mẫu, không chỉ tên — để trả lời được 'mẫu X có điều khoản gì'."""
    lark = FakeLarkKB(docx_with("Điều 1. Phạm vi. Điều 2. Giá trị {{gia_tri}}. " * 4))
    store, pf, eng, g, b = make(tmp_path, lark=lark)
    add_template(store)
    items = []
    monkeypatch.setattr(pf, "add_brain_item",
                        lambda title, content, status="approved", source_url=None:
                        items.append((title, content, source_url)), raising=False)
    assert flows.index_templates(b, log=lambda m: None) == 1
    title, content, url = items[0]
    assert title.startswith("[Mẫu hợp đồng]")
    assert "Điều 1. Phạm vi" in content          # nội dung thật
    assert "ten_ben_a" in content                # kèm danh sách field
    # chạy lần 2: mẫu không đổi → không index lại, khỏi rác brain
    items.clear()
    assert flows.index_templates(b, log=lambda m: None) == 0
    assert items == []


# ==================== S4 hằng tuần: nguồn theo nước + lưu Drive + index ====================

class FakeDrive(FakeLarkKB):
    """Drive giả có ensure_folder — kiểm việc tách folder theo nước."""

    def __init__(self):
        super().__init__(b"")
        self.folders, self.existing = {}, []

    def drive_files(self, folder):
        return self.existing

    def ensure_folder(self, parent, name):
        tok = self.folders.get(name)
        if not tok:
            tok = self.folders[name] = f"fld_{name}"
        return tok


def _add_item(store, key, country, url, doc_no=None, title="Văn bản"):
    store.write("INSERT INTO legal_news_items (key, country, doc_no, title, url, status, "
                "found_at) VALUES (?,?,?,?,?,'new',0)", (key, country, doc_no, title, url))


def test_sources_carry_country_and_dead_ones_start_inactive(tmp_path):
    """Seed phải phản ánh trạng thái ĐÃ KIỂM: nguồn chết để inactive kèm note, không
    để active rồi mỗi tuần báo lỗi mà không ai biết vì sao."""
    store, pf, eng, g, b = make(tmp_path)
    news.seed_sources(store)
    rows = news.sources(store, only_active=False)
    by_cc = {}
    for r in rows:
        by_cc.setdefault(r["country"], []).append(r)
    # Chốt 21/08: seed CHỈ có VN — Thái Lan bỏ khỏi scope, thêm trên console sau.
    assert set(by_cc) == {"VN"}
    active = [r for r in rows if r["active"]]
    # 3 nguồn VN đã kiểm chạy được: RSS LuatVietnam, thuvienphapluat, chinhphu
    assert {r["kind"] for r in active} == {"rss", "tvpl", "chinhphu"}
    for r in rows:
        if not r["active"]:
            assert r["note"], f"{r['name']} tắt mà không nói vì sao"


def test_seed_does_not_reactivate_source_admin_turned_off(tmp_path):
    store, pf, eng, g, b = make(tmp_path)
    news.seed_sources(store)
    store.write("UPDATE legal_news_sources SET active=0")
    news.seed_sources(store)
    assert not news.sources(store)      # vẫn tắt — seed không bật lại sau lưng admin


def test_html_source_without_pattern_reports_instead_of_silently_empty(tmp_path, monkeypatch):
    store, pf, eng, g, b = make(tmp_path)
    store.write("INSERT INTO legal_news_sources (name,url,kind,country,active) "
                "VALUES ('TH x','https://th.example/law','html','TH',1)")
    monkeypatch.setattr(news, "fetch", lambda *a, **k: "<a href='/law/1'>Thông tư</a>")
    assert news.crawl(store, log=lambda m: None) == []
    err = store.one("SELECT last_error FROM legal_news_sources WHERE name='TH x'")
    assert "link_pattern" in err["last_error"]


def test_html_source_with_pattern_extracts_docs(tmp_path, monkeypatch):
    store, pf, eng, g, b = make(tmp_path)
    store.write("INSERT INTO legal_news_sources (name,url,kind,country,link_pattern,active)"
                " VALUES ('TH gazette','https://th.example/list','html','TH','/law/',1)")
    monkeypatch.setattr(news, "fetch", lambda *a, **k: (
        '<a href="/law/123">ประกาศ 12/2569</a>'
        '<a href="/about">Về chúng tôi</a>'          # không khớp pattern → bỏ
        '<a href="/law/124">ประกาศ 13/2569</a>'))
    keys = news.crawl(store, log=lambda m: None)
    assert len(keys) == 2
    rows = store.query("SELECT * FROM legal_news_items")
    assert all(r["country"] == "TH" for r in rows)
    assert all(r["url"].startswith("https://th.example/law/") for r in rows)


def test_archive_saves_original_into_per_country_folder(tmp_path, monkeypatch):
    """Bản gốc về Drive, tách folder theo nước — không cần ai duyệt vì đó là tài liệu
    nhà nước, không phải nội dung AI sinh ra."""
    store, pf, eng, g, b = make(tmp_path, lark=FakeDrive())
    _add_item(store, "k-vn", "VN", "https://x.vn/nd-15.pdf", "15/2026/NĐ-CP")
    _add_item(store, "k-th", "TH", "https://x.th/notice.pdf", None, "ประกาศ")
    monkeypatch.setattr(news, "fetch_bytes", lambda *a, **k: b"%PDF-1.4 noi dung")
    assert news.archive(store, b.lark, ["k-vn", "k-th"], "fld_root",
                        log=lambda m: None) == 2
    assert set(b.lark.folders) == {"VN", "TH"}
    assert [f for f, _n, _d in b.lark.uploaded] == ["fld_VN", "fld_TH"]
    vn = store.one("SELECT * FROM legal_news_items WHERE key='k-vn'")
    assert vn["drive_url"].startswith("https://tenant/file/") and vn["status"] == "archived"


def test_archive_is_idempotent_and_survives_one_bad_download(tmp_path, monkeypatch):
    store, pf, eng, g, b = make(tmp_path, lark=FakeDrive())
    _add_item(store, "ok", "VN", "https://x.vn/a.pdf")
    _add_item(store, "bad", "VN", "https://x.vn/b.pdf")

    def fetch_bytes(url, timeout=60):
        if "b.pdf" in url:
            raise RuntimeError("404")
        return b"data"

    monkeypatch.setattr(news, "fetch_bytes", fetch_bytes)
    assert news.archive(store, b.lark, ["ok", "bad"], "fld_root", log=lambda m: None) == 1
    # chạy lại: mục đã có drive_url thì không tải/upload lần nữa
    assert news.archive(store, b.lark, ["ok", "bad"], "fld_root", log=lambda m: None) == 0


def test_archive_without_folder_config_is_skipped_not_crash(tmp_path):
    store, pf, eng, g, b = make(tmp_path, lark=FakeDrive())
    _add_item(store, "k", "VN", "https://x.vn/a.pdf")
    assert news.archive(store, b.lark, ["k"], None, log=lambda m: None) == 0


def test_index_points_to_where_to_retrieve(tmp_path):
    """Mục index phải nói rõ: nước, số hiệu, link nguồn gốc VÀ link bản lưu nội bộ."""
    store, pf, eng, g, b = make(tmp_path)
    _add_item(store, "k1", "TH", "https://th.example/law/1", "12/2569", "ประกาศ นาฬิกา")
    store.write("UPDATE legal_news_items SET drive_url='https://tenant/file/tok1' "
                "WHERE key='k1'")
    assert flows.index_legal_docs(b, ["k1"], log=lambda m: None) == 1
    title, content, url = pf.brain_items[0]
    assert "[Văn bản pháp luật · TH]" in title and "12/2569" in title
    assert "Nguồn gốc: https://th.example/law/1" in content
    assert "Lark Drive): https://tenant/file/tok1" in content
    assert url == "https://tenant/file/tok1"


def test_index_says_when_no_internal_copy(tmp_path):
    store, pf, eng, g, b = make(tmp_path)
    _add_item(store, "k2", "VN", "https://x.vn/a")
    flows.index_legal_docs(b, ["k2"], log=lambda m: None)
    _t, content, url = pf.brain_items[0]
    assert "CHƯA có" in content and url == "https://x.vn/a"


def test_index_skips_item_without_source_url(tmp_path):
    store, pf, eng, g, b = make(tmp_path)
    store.write("INSERT INTO legal_news_items (key, country, title, url, status, found_at) "
                "VALUES ('k3','VN','Không link','','new',0)")
    assert flows.index_legal_docs(b, ["k3"], log=lambda m: None) == 0
    assert pf.brain_items == []


def test_weekly_cycle_archives_and_indexes_before_gate(tmp_path, monkeypatch):
    """Lưu + index xảy ra TRƯỚC gate: dữ kiện có ngay, chỉ phần model diễn giải mới chờ duyệt."""
    store, pf, eng, g, b = make(tmp_path, lark=FakeDrive())
    store.write("INSERT INTO legal_news_sources (name,url,kind,country,active) "
                "VALUES ('LVN','https://x.vn/rss','rss','VN',1)")
    monkeypatch.setattr(news, "fetch", lambda *a, **k:
                        "<rss><channel><item><title>Nghị định 21/2026/NĐ-CP</title>"
                        "<link>https://x.vn/nd21.pdf</link></item></channel></rss>")
    monkeypatch.setattr(news, "fetch_bytes", lambda *a, **k: b"pdf")
    monkeypatch.setattr(flows.brain, "call_claude", lambda *a, **k: json.dumps(
        {"scope": "s", "effective": "e", "impact": "i"}))
    monkeypatch.setenv("LEGAL_DRIVE_FOLDER", "fld_root")
    gid = flows.news_cycle(b, GROUP, log=lambda m: None)
    it = store.one("SELECT * FROM legal_news_items WHERE doc_no='21/2026/NĐ-CP'")
    assert it["drive_url"], "phải lưu bản gốc trước khi mở gate"
    assert pf.brain_items, "phải index trước khi mở gate"
    assert b.gates.get(gid)["status"] == "open"      # digest vẫn chờ Pháp chế duyệt


# ==================== thuvienphapluat.vn (adapter tvpl) ====================

SEARCH_HTML = """
<div class="number"> 1</div><div class="nq">
  <p class="nqTitle" lawid='720746'>
    <a onclick="Doc_CT(MemberGA)" href="https://thuvienphapluat.vn/cong-van/Thu-tuc/Cong-van-529-TANDTC-PC-2026-huong-dan-720746.aspx">Công văn 529/TANDTC-PC năm 2026 hướng dẫn</a>
  </p></div>
<div class="nq"><p class="nqTitle" lawid='720558'>
    <a href="/cong-van/Doanh-nghiep/Cong-van-20519-CHQ-GSQL-2026-thu-tuc-720558.aspx">Công văn 20519/CHQ-GSQL năm 2026 thủ tục</a>
  </p></div>
<p class="nqTitle" lawid='999'>
  <a href="https://thuvienphapluat.vn/van-ban/X/Nghi-quyet-148-NQ-CP-579993.aspx?v=tvpl-hdsd-firsr&step=step6">Theo dõi hiệu lực văn bản;</a>
</p>
"""

DOC_HTML = ('<html><body><div id="divContentDoc"><p>CHÍNH PHỦ</p>'
            '<p>Số: 326/2026/NĐ-CP</p><p>NGHỊ ĐỊNH QUY ĐỊNH VỀ ĐỊNH DANH</p>'
            '</div><script>x</script></body></html>')


def test_tvpl_search_skips_help_links(monkeypatch):
    """Link hướng dẫn (v=tvpl-hdsd…&step=) cũng nằm dưới /van-ban/ — lọc theo đường dẫn
    thì lẫn vào và thành "tìm được văn bản" mà thực chất là trang trợ giúp."""
    from legalkb import tvpl
    monkeypatch.setattr(tvpl, "get", lambda *a, **k: SEARCH_HTML)
    docs = tvpl.search(log=lambda m: None)
    assert [d["lawid"] for d in docs] == ["720746", "720558"]
    assert all("hdsd" not in d["url"] for d in docs)
    assert docs[1]["url"].startswith("https://thuvienphapluat.vn/cong-van/")  # url tuyệt đối


def test_tvpl_extracts_full_text(monkeypatch):
    from legalkb import tvpl
    monkeypatch.setattr(tvpl, "get", lambda *a, **k: DOC_HTML)
    t = tvpl.fetch_text("https://x/1.aspx")
    assert "Số: 326/2026/NĐ-CP" in t and "ĐỊNH DANH" in t
    assert "<p>" not in t and "script" not in t


def test_tvpl_download_original_needs_cookie_not_password(monkeypatch):
    """Agent KHÔNG giữ mật khẩu của ai: tải file gốc chỉ chạy khi có cookie phiên do
    người dùng tự lấy. Toàn văn đã lấy được nên đây chỉ là tuỳ chọn."""
    from legalkb import tvpl
    monkeypatch.delenv("TVPL_COOKIE", raising=False)
    try:
        tvpl.download_original("https://x/1.aspx")
    except RuntimeError as e:
        assert "TVPL_COOKIE" in str(e) and "fetch_text" in str(e)
    else:
        raise AssertionError("thiếu cookie thì phải báo lỗi rõ")
    src = (pathlib.Path(__file__).resolve().parent.parent / "legalkb" / "tvpl.py").read_text()
    assert "TVPL_PASSWORD" not in src and "password" not in src.lower().replace(
        "mật khẩu", ""), "không được nhận mật khẩu"


def test_tvpl_crawl_and_archive_saves_extracted_text(tmp_path, monkeypatch):
    """Nguồn tvpl: lưu TOÀN VĂN .txt, không lưu HTML 400KB toàn menu/quảng cáo."""
    from legalkb import tvpl
    store, pf, eng, g, b = make(tmp_path, lark=FakeDrive())
    store.write("INSERT INTO legal_news_sources (name,url,kind,country,active) VALUES "
                "('TVPL','https://thuvienphapluat.vn/page/tim-van-ban.aspx','tvpl','VN',1)")
    monkeypatch.setattr(tvpl, "get", lambda *a, **k: SEARCH_HTML)
    keys = news.crawl(store, log=lambda m: None)
    assert set(keys) == {"529/TANDTC-PC", "20519/CHQ-GSQL"}   # dedupe theo số hiệu

    monkeypatch.setattr(tvpl, "fetch_text", lambda url, cookie=None: "TOÀN VĂN NGHỊ ĐỊNH")
    assert news.archive(store, b.lark, keys, "fld_root", log=lambda m: None) == 2
    names = [n for _f, n, _d in b.lark.uploaded]
    assert all(n.endswith(".txt") for n in names), names
    assert any("529/TANDTC-PC".replace("/", " ") in n or "529" in n for n in names)
    assert all(d == "TOÀN VĂN NGHỊ ĐỊNH".encode() for _f, _n, d in b.lark.uploaded)


def test_tvpl_archive_reports_when_layout_changed(tmp_path, monkeypatch):
    from legalkb import tvpl
    store, pf, eng, g, b = make(tmp_path, lark=FakeDrive())
    store.write("INSERT INTO legal_news_sources (name,url,kind,country,active) VALUES "
                "('TVPL','https://x/search','tvpl','VN',1)")
    sid = store.one("SELECT id FROM legal_news_sources")["id"]
    store.write("INSERT INTO legal_news_items (key, source_id, country, title, url, status,"
                " found_at) VALUES ('k',?,'VN','T',"
                "'https://thuvienphapluat.vn/van-ban/1.aspx','new',0)", (sid,))
    monkeypatch.setattr(tvpl, "fetch_text", lambda url, cookie=None: "")
    logs = []
    assert news.archive(store, b.lark, ["k"], "fld_root", log=logs.append) == 0
    assert any("không trích được toàn văn" in m for m in logs)


def test_doc_no_handles_five_digit_numbers():
    """20519/CHQ-GSQL — công văn hải quan 5 chữ số. Giới hạn \\d{1,4} cũ bỏ sót thật."""
    assert news.doc_no_of("Công văn 20519/CHQ-GSQL năm 2026") == "20519/CHQ-GSQL"
    assert news.doc_no_of("Nghị định 326/2026/NĐ-CP") == "326/2026/NĐ-CP"


# ==================== chinhphu.vn (adapter chinhphu) ====================

# Dòng đầu là dòng TIÊU ĐỀ BẢNG — cũng khớp <tr><td> nhưng không có span.code.
CP_HTML = """
<tr><td><b>Số hiệu</b></td><td><b>Ngày</b></td><td><b>Trích yếu</b></td></tr>
<tr>
  <td><a href='/?pageid=27160&docid=219221'><span class="code">326/2026/NĐ-CP</span></a></td>
  <td><span class="issued-date">19/08/2026</span></td>
  <td><a href='/?pageid=27160&docid=219221'><span class="substract">Quy định về định danh địa điểm
      </span></a>
    <div class="bl-doc-files">
      <div class="bl-doc-file"><a href="https://datafiles.chinhphu.vn/cpp/files/vbpq/2026/8/326_2026_nd-cp-signed.pdf" download>Tài liệu đính kèm</a></div>
      <div class="bl-doc-file"><a href="https://datafiles.chinhphu.vn/cpp/files/vbpq/2026/8/326_2026_nd-cp_pl.pdf" download>Tài liệu đính kèm</a></div>
    </div>
  </td>
</tr>
<tr>
  <td><a href='/?pageid=27160&docid=219231'><span class="code">1602/QĐ-TTg</span></a></td>
  <td><span class="issued-date">19/08/2026</span></td>
  <td><a href='/?pageid=27160&docid=219231'><span class="substract">Về việc giải thể Ban chỉ đạo</span></a>
    <div class="bl-doc-file"><a href="https://datafiles.chinhphu.vn/cpp/files/vbpq/2026/8/1602_qd-ttg-signed.pdf">Tài liệu đính kèm</a></div>
  </td>
</tr>
"""


def test_chinhphu_takes_doc_no_from_source_column(monkeypatch):
    """Số hiệu là DỮ LIỆU của nguồn, không phải regex dò tiêu đề.

    Trích yếu "Quy định về định danh địa điểm" không có một chữ số nào — dò regex là trượt.
    """
    from legalkb import chinhphu
    monkeypatch.setattr(chinhphu.web, "get", lambda *a, **k: CP_HTML)
    docs = chinhphu.search(log=lambda m: None)
    assert [d["doc_no"] for d in docs] == ["326/2026/NĐ-CP", "1602/QĐ-TTg"]
    assert news.doc_no_of(docs[0]["title"]) is None, "tiêu đề không có số → phải lấy từ cột"
    assert docs[0]["title"] == "Quy định về định danh địa điểm"
    assert docs[0]["issued"] == "19/08/2026"
    assert len(docs[0]["files"]) == 2 and docs[0]["files"][0].endswith("-signed.pdf")
    assert docs[0]["url"] == "https://chinhphu.vn/?pageid=27160&docid=219221"


def test_chinhphu_rejects_error_page_pretending_to_be_pdf(monkeypatch):
    """Nguồn trả trang lỗi thay vì PDF → phải báo lỗi, không lưu 5KB HTML thành "văn bản"."""
    from legalkb import chinhphu
    monkeypatch.setattr(chinhphu.web, "get_bytes", lambda *a, **k: b"<html>404</html>")
    with pytest.raises(RuntimeError, match="không phải PDF"):
        chinhphu.download_original("https://chinhphu.vn/?docid=1",
                                   ["https://datafiles.chinhphu.vn/x.pdf"])


def test_chinhphu_no_attachment_is_reported_not_silently_skipped(monkeypatch):
    from legalkb import chinhphu
    monkeypatch.setattr(chinhphu.web, "get", lambda *a, **k: "<html>không có file</html>")
    with pytest.raises(RuntimeError, match="không có file đính kèm"):
        chinhphu.download_original("https://chinhphu.vn/?docid=1")


def test_crawl_chinhphu_keys_by_doc_no_and_archives_signed_pdf(tmp_path, monkeypatch):
    from legalkb import chinhphu
    store, pf, eng, g, b = make(tmp_path, lark=FakeDrive())
    store.write("INSERT INTO legal_news_sources (name,url,kind,country,active) VALUES "
                "('CP','https://chinhphu.vn/he-thong-van-ban','chinhphu','VN',1)")
    monkeypatch.setattr(chinhphu.web, "get", lambda *a, **k: CP_HTML)
    keys = news.crawl(store, log=lambda m: None)
    assert set(keys) == {"326/2026/NĐ-CP", "1602/QĐ-TTg"}

    monkeypatch.setattr(chinhphu.web, "get_bytes", lambda *a, **k: b"%PDF-1.7 noi dung")
    assert news.archive(store, b.lark, keys, "fld_root", log=lambda m: None) == 2
    names = [n for _f, n, _d in b.lark.uploaded]
    assert all(n.endswith(".pdf") for n in names), names
    assert all(d.startswith(b"%PDF") for _f, _n, d in b.lark.uploaded)
    # file_urls lưu lúc crawl → archive không phải đọc lại trang chi tiết
    it = store.one("SELECT file_urls FROM legal_news_items WHERE key='326/2026/NĐ-CP'")
    assert len(json.loads(it["file_urls"])) == 2


def test_same_document_from_two_sources_stored_once(tmp_path, monkeypatch):
    """Dedupe LIÊN NGUỒN chỉ chạy được nhờ số hiệu — đây là lý do phải lấy số hiệu thật."""
    from legalkb import chinhphu
    store, pf, eng, g, b = make(tmp_path, lark=FakeDrive())
    store.write("INSERT INTO legal_news_sources (name,url,kind,country,active) VALUES "
                "('CP','https://chinhphu.vn/he-thong-van-ban','chinhphu','VN',1)")
    monkeypatch.setattr(chinhphu.web, "get", lambda *a, **k: CP_HTML)
    assert len(news.crawl(store, log=lambda m: None)) == 2

    store.write("INSERT INTO legal_news_sources (name,url,kind,country,active) VALUES "
                "('RSS','https://luatvietnam.vn/rss/x.rss','rss','VN',1)")
    rss = ("<rss><item><title>Nghị định 326/2026/NĐ-CP về định danh địa điểm</title>"
           "<link>https://luatvietnam.vn/abc-1-d1.html</link></item></rss>")
    monkeypatch.setattr(news, "fetch", lambda *a, **k: rss)
    assert news.crawl(store, log=lambda m: None) == [], "cùng số hiệu → không lưu lần hai"
    assert store.one("SELECT COUNT(*) c FROM legal_news_items")["c"] == 2


# ==================== luatvietnam.vn (tra cứu theo tên) ====================

LVN_SEARCH = """
<a href="https://luatvietnam.vn/lao-dong/bo-luat-lao-dong-2019-so-45-2019-qh14-179015-d1.html">Bộ luật Lao động 2019, số 45/2019/QH14</a>
<a href="https://luatvietnam.vn/lao-dong/du-thao-bo-luat-lao-dong-sua-doi-176972-d10.html">Dự thảo Bộ luật Lao động sửa đổi</a>
<a href="https://luatvietnam.vn/tin-phap-luat/diem-tin-van-ban-moi-123.html">Điểm tin văn bản mới</a>
"""
LVN_DOC = """<div class="the-document-body">
<p>Điều 1. Phạm vi điều chỉnh</p><p>Bộ luật Lao động quy định tiêu chuẩn lao động.</p>
<p>Đang theo dõi</p><p>Điều 2. Đối tượng áp dụng</p><p>Đang theo dõi</p><p>1. Người lao động.</p>
<footer>chân trang</footer>"""


def test_luatvietnam_excludes_drafts_by_default(monkeypatch):
    """72/123 kết quả trên trang tìm kiếm là DỰ THẢO — trích dự thảo như đang có hiệu lực
    là sai nghiêm trọng với việc pháp chế."""
    from legalkb import luatvietnam as lvn
    monkeypatch.setattr(lvn.web, "get", lambda *a, **k: LVN_SEARCH)
    docs = lvn.search("bo luat lao dong", log=lambda m: None)
    assert [d["title"] for d in docs] == ["Bộ luật Lao động 2019, số 45/2019/QH14"]
    assert docs[0]["is_draft"] is False
    withdraft = lvn.search("x", include_drafts=True, log=lambda m: None)
    assert [d["is_draft"] for d in withdraft] == [False, True]
    assert all("tin-phap-luat" not in d["url"] for d in withdraft), "bài tin không phải văn bản"


def test_luatvietnam_strips_ui_buttons_from_text(monkeypatch):
    """"Đang theo dõi" là nút bấm của giao diện, chèn sau 148/218 điều — không phải nội dung."""
    from legalkb import luatvietnam as lvn
    monkeypatch.setattr(lvn.web, "get", lambda *a, **k: LVN_DOC)
    t = lvn.fetch_text("https://luatvietnam.vn/x-1-d1.html")
    assert "Đang theo dõi" not in t
    assert "Điều 1. Phạm vi điều chỉnh" in t and "1. Người lao động." in t
    assert "chân trang" not in t


def test_luatvietnam_layout_change_returns_empty_not_garbage(monkeypatch):
    from legalkb import luatvietnam as lvn
    monkeypatch.setattr(lvn.web, "get", lambda *a, **k: "<html><body>đổi layout</body></html>")
    assert lvn.fetch_text("https://luatvietnam.vn/x-1-d1.html") == ""


# ==================== S4 hỏi đáp: tra cứu tức thời ====================

def test_s4_answer_looks_up_the_document_the_question_names(tmp_path, monkeypatch):
    """Trước đây hàm này bỏ qua câu hỏi và luôn trả 5 văn bản mới nhất."""
    from legalkb import luatvietnam as lvn
    store, pf, eng, g, b = make(tmp_path, lark=FakeDrive())
    news.seed_sources(store)
    monkeypatch.setattr(flows.brain, "call_claude", lambda *a, **k: "Bộ luật Lao động 2019")
    monkeypatch.setattr(lvn.web, "get", lambda *a, **k: LVN_SEARCH)
    monkeypatch.setattr(lvn, "fetch_text", lambda url: "TOÀN VĂN BỘ LUẬT")
    monkeypatch.setenv(flows.LAW_ARCHIVE_ENV, "fld_law")
    out = flows.s4_answer(b, "Bộ luật Lao động quy định gì về thử việc?")
    assert "Bộ luật Lao động 2019" in out and "179015-d1.html" in out
    assert "bản lưu nội bộ" in out, "phải chỉ được cả bản lưu trong Drive"
    assert b.lark.uploaded and b.lark.uploaded[0][2] == "TOÀN VĂN BỘ LUẬT".encode()


def test_s4_answer_uses_doc_no_in_question_without_calling_model(tmp_path, monkeypatch):
    store, pf, eng, g, b = make(tmp_path, lark=FakeDrive())
    news.seed_sources(store)

    def boom(*a, **k):
        raise AssertionError("câu hỏi đã có số hiệu → không cần tốn lượt model")
    monkeypatch.setattr(flows.brain, "call_claude", boom)
    monkeypatch.setattr(flows, "lookup_law", lambda b_, kw, log=print: [])
    out = flows.s4_answer(b, "Nghị định 98/2020/NĐ-CP xử phạt thế nào?")
    assert "98/2020/NĐ-CP" in out


def test_s4_answer_refuses_to_guess_when_nothing_found(tmp_path, monkeypatch):
    store, pf, eng, g, b = make(tmp_path, lark=FakeDrive())
    news.seed_sources(store)
    monkeypatch.setattr(flows.brain, "call_claude", lambda *a, **k: "Luật Không Tồn Tại")
    monkeypatch.setattr(flows, "lookup_law", lambda b_, kw, log=print: [])
    out = flows.s4_answer(b, "Luật Không Tồn Tại nói gì?")
    assert "không ra văn bản nào" in out and "không\nđoán".replace("\n", " ") in out


def test_s4_answer_marks_drafts_loudly(tmp_path):
    store, pf, eng, g, b = make(tmp_path, lark=FakeDrive())
    store.write("INSERT INTO legal_news_items (key,country,doc_no,title,url,is_draft,status,"
                "found_at) VALUES ('k','VN',NULL,'Dự thảo Luật X','https://x/1-d10.html',"
                "1,'published',1)")
    out = flows.s4_answer(b, "có gì mới không")
    assert "DỰ THẢO" in out and "chưa ban hành" in out


def test_lookup_law_survives_source_being_down(tmp_path, monkeypatch):
    """Nguồn ngoài chết thì trả rỗng — không được làm vỡ lượt trả lời."""
    from legalkb import luatvietnam as lvn
    store, pf, eng, g, b = make(tmp_path, lark=FakeDrive())
    news.seed_sources(store)

    def down(*a, **k):
        raise OSError("timeout")
    monkeypatch.setattr(lvn.web, "get", down)
    assert flows.lookup_law(b, "Bộ luật Lao động", log=lambda m: None) == []


# ==================== web.py: lịch sự với nguồn ====================

def test_web_gap_is_per_host_not_global(monkeypatch):
    """Ba nguồn khác nhau không có lý gì phải chờ nhau; hai lời gọi CÙNG nguồn thì phải."""
    from legalkb import web
    slept, now = [], [1000.0]
    monkeypatch.setattr(web.time, "sleep", lambda s: slept.append(round(s, 2)))
    monkeypatch.setattr(web.time, "time", lambda: now[0])
    web._last.clear()
    web._wait("https://a.vn/1", 2)
    web._wait("https://b.vn/1", 2)
    assert slept == [], "host khác nhau → không chờ"
    web._wait("https://a.vn/2", 2)
    assert slept == [2.0], "cùng host → chờ đủ giãn cách"


def test_model_junk_is_not_treated_as_a_document_name(tmp_path, monkeypatch):
    """Lỗi THẬT do test bắt được: model trả khối JSON, bản trước đem đi tra và trích dẫn
    3 văn bản không liên quan như thể đã tra đúng."""
    store, pf, eng, g, b = make(tmp_path, lark=FakeDrive())
    monkeypatch.setattr(flows.brain, "call_claude",
                        lambda *a, **k: '{"intent": "s1_qa", "risk": "low"}')
    monkeypatch.setattr(flows, "lookup_law",
                        lambda *a, **k: pytest.fail("không được tra bằng rác của model"))
    assert flows._doc_name_in(b, "có gì mới không") is None
    assert "chưa có bản tin" in flows.s4_answer(b, "có gì mới không")


def test_lookup_rejects_hits_that_do_not_match_the_name(tmp_path, monkeypatch):
    """luatvietnam trả kết quả cho gần như mọi truy vấn → "có kết quả" không phải bằng
    chứng đã tra đúng văn bản."""
    from legalkb import luatvietnam as lvn
    store, pf, eng, g, b = make(tmp_path, lark=FakeDrive())
    news.seed_sources(store)
    off_topic = ('<a href="https://luatvietnam.vn/tieu-chuan/tcvn-iso-31000-2011-quan-ly-'
                 'rui-ro-185124-d3.html">TCVN ISO 31000:2011 quản lý rủi ro</a>')
    monkeypatch.setattr(lvn.web, "get", lambda *a, **k: off_topic)
    logs = []
    assert flows.lookup_law(b, "Luật Không Tồn Tại", log=logs.append) == []
    assert any("không khớp tên" in m for m in logs)
    assert store.one("SELECT COUNT(*) c FROM legal_news_items")["c"] == 0


def test_doc_no_keeps_mixed_case_and_numeric_suffixes():
    """Ba lỗi chạy thật: 55/CĐ-TTg bị cắt còn 55/CĐ (⇒ lưu hai bản), và 45/2019/QH14
    hỏng luôn cả cụm ⇒ Bộ luật Lao động thành "không có số hiệu"."""
    assert news.doc_no_of("Công điện 55/CĐ-TTg của Thủ tướng") == "55/CĐ-TTg"
    assert news.doc_no_of("Bộ luật Lao động 2019, số 45/2019/QH14") == "45/2019/QH14"
    assert news.doc_no_of("Quyết định 3523/QĐ-BKHCN của Bộ") == "3523/QĐ-BKHCN"
    assert news.doc_no_of("không có số hiệu nào ở đây") is None


def test_better_original_replaces_weaker_one_found_first(tmp_path, monkeypatch):
    """RSS chạy trước và chỉ có link trang tin; chinhphu.vn tới sau nhưng có PDF ký số.
    Dedupe giữ bản thấy trước, nhưng bản GỐC phải là bản tốt hơn."""
    from legalkb import chinhphu
    store, pf, eng, g, b = make(tmp_path, lark=FakeDrive())
    store.write("INSERT INTO legal_news_sources (name,url,kind,country,active) VALUES "
                "('RSS','https://luatvietnam.vn/rss/x.rss','rss','VN',1)")
    rss = ("<rss><item><title>Nghị định 326/2026/NĐ-CP về định danh địa điểm</title>"
           "<link>https://luatvietnam.vn/abc-1-d1.html</link></item></rss>")
    monkeypatch.setattr(news, "fetch", lambda *a, **k: rss)
    assert news.crawl(store, log=lambda m: None) == ["326/2026/NĐ-CP"]
    assert store.one("SELECT file_urls FROM legal_news_items")["file_urls"] is None

    store.write("INSERT INTO legal_news_sources (name,url,kind,country,active) VALUES "
                "('CP','https://chinhphu.vn/he-thong-van-ban','chinhphu','VN',1)")
    monkeypatch.setattr(chinhphu.web, "get", lambda *a, **k: CP_HTML)
    logs = []
    news.crawl(store, log=logs.append)
    it = store.one("SELECT url, file_urls FROM legal_news_items WHERE key='326/2026/NĐ-CP'")
    assert "chinhphu.vn" in it["url"] and "signed.pdf" in it["file_urls"]
    assert any("đổi sang bản gốc" in m for m in logs)

    monkeypatch.setattr(chinhphu.web, "get_bytes", lambda *a, **k: b"%PDF-1.7 x")
    news.archive(store, b.lark, ["326/2026/NĐ-CP"], "fld", log=lambda m: None)
    assert b.lark.uploaded[0][1].endswith(".pdf")


def test_never_saves_whole_page_html_as_the_original(tmp_path, monkeypatch):
    """Trang nguồn 2MB toàn menu/quảng cáo — lưu nó rồi gọi là "văn bản gốc" thì người
    sau mở ra không đọc được gì, mà hệ thống báo đã lưu thành công."""
    store, pf, eng, g, b = make(tmp_path, lark=FakeDrive())
    store.write("INSERT INTO legal_news_items (key,country,title,url,status,found_at) "
                "VALUES ('k','VN','T','https://nguon-la.vn/tin/abc','new',0)")
    monkeypatch.setattr(news, "fetch_bytes", lambda *a, **k: b"<html>" + b"x" * 5000)
    logs = []
    assert news.archive(store, b.lark, ["k"], "fld", log=logs.append) == 0
    assert not b.lark.uploaded
    assert any("chưa có bộ trích toàn văn" in m for m in logs)


@pytest.mark.parametrize("keyword,title,ok", [
    ("Bộ luật Lao động 2019", "Bộ luật Lao động 2019 , số 45/ 2019 /QH14", True),
    # Lỗi THẬT lộ khi chạy live: đếm tỉ lệ từ trùng thì văn bản này đạt 3/4
    # (luật · động · 2019) và được lưu + index như thể là Bộ luật Lao động.
    ("Bộ luật Lao động 2019", "Luật Lực lượng dự bị động viên 2019 , số 53/2019/QH14", False),
    ("luật lao động", "Bộ luật Lao động 2019, số 45/2019/QH14", True),
    # Bỏ từ ở GIỮA làm cụm mất liền mạch → chỉ bỏ năm ở CUỐI tên
    ("Luật Bảo vệ quyền lợi người tiêu dùng",
     "Luật Bảo vệ quyền lợi người tiêu dùng 2023", True),
    # Hỏi bằng số hiệu: nguồn hay chèn khoảng trắng "98/ 2020 /NĐ-CP"
    ("98/2020/NĐ-CP", "Nghị định 98/ 2020 /NĐ-CP xử phạt", True),
    ("98/2020/NĐ-CP", "Nghị định 326/2026/NĐ-CP định danh địa điểm", False),
    ("Luật Không Tồn Tại", "TCVN ISO 31000:2011 quản lý rủi ro", False),
])
def test_lookup_matches_by_phrase_not_by_loose_word_overlap(keyword, title, ok):
    assert flows._matches(keyword, title) is ok


# ==================== nạp tay: người bỏ file vào kho Drive ====================

class FakeDriveTree:
    """Drive giả có cây folder: root → VN/, TH/ → file."""

    def __init__(self, tree):
        self.tree = tree      # {folder_token: [entry, ...]}
        self.uploaded = []

    def drive_files(self, token):
        return self.tree.get(token, [])

    def drive_file_url(self, token, file_type="file"):
        return f"https://tenant/file/{token}"


def _tree():
    return {
        "root": [{"token": "fVN", "name": "VN", "type": "folder"},
                 {"token": "fTH", "name": "TH", "type": "folder"}],
        "fVN": [{"token": "v1", "name": "Nghị định 99/2026/NĐ-CP về nhãn hàng hoá.pdf",
                 "type": "file"}],
        "fTH": [{"token": "t1", "name": "Notification MOC labelling 2026.pdf", "type": "file"},
                {"token": "t2", "name": "Royal Gazette vol 143 retail.pdf", "type": "file"}],
    }


def test_manual_drive_files_become_indexable(tmp_path):
    """Thái Lan không crawl tự động được (9 nguồn đã đo) → người bỏ tay vào folder TH/
    thì vẫn phải vào index, nếu không thì kho có văn bản mà agent không biết."""
    store, pf, eng, g, b = make(tmp_path, lark=FakeDriveTree(_tree()))
    keys = news.ingest_drive_folder(store, b.lark, "root", log=lambda m: None)
    assert len(keys) == 3
    rows = {r["country"] for r in store.query("SELECT country FROM legal_news_items")}
    assert rows == {"VN", "TH"}
    th = store.query("SELECT * FROM legal_news_items WHERE country='TH'")
    assert len(th) == 2
    assert all(r["drive_url"] and r["status"] == "archived" for r in th)


def test_manual_ingest_is_idempotent(tmp_path):
    store, pf, eng, g, b = make(tmp_path, lark=FakeDriveTree(_tree()))
    news.ingest_drive_folder(store, b.lark, "root", log=lambda m: None)
    again = news.ingest_drive_folder(store, b.lark, "root", log=lambda m: None)
    assert again == [], "quét lại không được tạo bản ghi trùng"


def test_manual_ingest_skips_doc_already_crawled(tmp_path):
    """Cùng số hiệu đã crawl về thì file bỏ tay KHÔNG tạo bản ghi thứ hai."""
    store, pf, eng, g, b = make(tmp_path, lark=FakeDriveTree(_tree()))
    store.write("INSERT INTO legal_news_items (key,country,doc_no,title,url,status,found_at)"
                " VALUES ('99/2026/NĐ-CP','VN','99/2026/NĐ-CP','đã crawl','https://x','new',1)")
    keys = news.ingest_drive_folder(store, b.lark, "root", log=lambda m: None)
    assert len(keys) == 2, "chỉ 2 file TH được nhận"
    assert store.one("SELECT COUNT(*) c FROM legal_news_items WHERE doc_no='99/2026/NĐ-CP'")["c"] == 1


def test_manual_ingest_survives_unreadable_folder(tmp_path):
    class Broken(FakeDriveTree):
        def drive_files(self, token):
            if token == "fTH":
                raise RuntimeError("1061004 forbidden")
            return super().drive_files(token)
    store, pf, eng, g, b = make(tmp_path, lark=Broken(_tree()))
    logs = []
    keys = news.ingest_drive_folder(store, b.lark, "root", log=logs.append)
    assert len(keys) == 1, "folder VN vẫn nhận được"
    assert any("không đọc được folder TH" in m for m in logs)


@pytest.mark.parametrize("filename,expected", [
    # `_safe_name()` làm phẳng dấu "/" khi upload → phải đọc lại được số hiệu
    ("326 2026 NĐ-CP Nghị định quy định về định danh địa điểm.pdf", "326/2026/NĐ-CP"),
    ("55 CĐ-TTg Công điện của Thủ tướng.pdf", "55/CĐ-TTg"),
    ("45 2019 QH14 Bộ luật Lao động 2019.txt", "45/2019/QH14"),
    ("5947 CT-KTr Công văn năm 2026.txt", "5947/CT-KTr"),
    # chốt chống nhận bừa: không có năm VÀ không có gạch nối → không phải số hiệu
    ("2026 BAO CAO nam.pdf", None),
    ("Notification MOC labelling 2026.pdf", None),
    ("Royal Gazette vol 143 retail.pdf", None),
    # tên còn nguyên dấu "/" thì đường cũ vẫn chạy
    ("Nghị định 326/2026/NĐ-CP.pdf", "326/2026/NĐ-CP"),
])
def test_doc_no_recovered_from_flattened_filename(filename, expected):
    """Lỗi thật: mọi văn bản nạp tay đều doc_no=NULL ⇒ dedupe không chạy ⇒ Bộ luật Lao
    động vào kho hai lần."""
    assert news.doc_no_from_filename(filename) == expected


def test_manual_ingest_dedupes_by_recovered_doc_no(tmp_path):
    tree = {"root": [{"token": "fVN", "name": "VN", "type": "folder"}],
            "fVN": [{"token": "a", "name": "45 2019 QH14 Bộ luật Lao động 2019.txt",
                     "type": "file"},
                    {"token": "b", "name": "45 2019 QH14 Bộ luật Lao động 2019.txt",
                     "type": "file"}]}
    store, pf, eng, g, b = make(tmp_path, lark=FakeDriveTree(tree))
    keys = news.ingest_drive_folder(store, b.lark, "root", log=lambda m: None)
    assert len(keys) == 1, "hai file cùng số hiệu → một bản ghi"
    assert store.one("SELECT doc_no FROM legal_news_items")["doc_no"] == "45/2019/QH14"


# ==================== mẫu hợp đồng: NHIỀU folder Drive ====================

class MultiFolderDrive:
    def __init__(self, folders):
        self.folders = folders
        self.downloads = []

    def drive_files(self, token):
        if token not in self.folders:
            raise RuntimeError(f"1061004 forbidden: {token}")
        return self.folders[token]

    def drive_download(self, token):
        self.downloads.append(token)
        return docx_with("Bên A: {{ten_ben_a}} — giá trị {{gia_tri}}")

    def drive_file_url(self, token, file_type="file"):
        return f"https://tenant/file/{token}"


def test_templates_scanned_from_every_configured_folder(tmp_path):
    """Chốt 21/08: mẫu ở Drive, và có HAI folder tên đều hợp lý. Đoán sai một cái là S2
    không thấy mẫu nào mà vẫn báo thành công."""
    store, pf, eng, g, b = make(tmp_path)
    drive = MultiFolderDrive({
        "fA": [{"token": "t1", "name": "[MAU]_HĐ thuê mặt bằng.docx", "type": "file"}],
        "fB": [{"token": "t2", "name": "[MAU]_HĐ cung cấp dịch vụ.docx", "type": "file"}],
    })
    r = contracts.sync_templates(drive, store, "fA,fB", log=lambda m: None)
    assert r["templates"] == 2
    names = {t["name"] for t in contracts.templates(store)}
    assert names == {"[MAU]_HĐ thuê mặt bằng", "[MAU]_HĐ cung cấp dịch vụ"}


def test_draft_subfolder_is_never_scanned_as_a_template(tmp_path):
    """Folder bản thảo agent tự xuất nằm TRONG folder mẫu. Quét đệ quy = bản thảo của
    chính mình thành mẫu cho lần sau, sai lệch tích luỹ mà không ai thấy."""
    store, pf, eng, g, b = make(tmp_path)
    drive = MultiFolderDrive({
        "fA": [{"token": "t1", "name": "[MAU]_HĐ mẫu thật.docx", "type": "file"},
               {"token": "fDraft", "name": "BAN THAO DRAFT (agent xuat)",
                "type": "folder"}],
        "fDraft": [{"token": "d9", "name": "[MAU]_HĐ ABC DRAFT.docx", "type": "file"}],
    })
    contracts.sync_templates(drive, store, "fA", log=lambda m: None)
    assert {t["name"] for t in contracts.templates(store)} == {"[MAU]_HĐ mẫu thật"}


def test_one_unreadable_folder_does_not_lose_the_other(tmp_path):
    store, pf, eng, g, b = make(tmp_path)
    drive = MultiFolderDrive({"fA": [{"token": "t1", "name": "[MAU]_HĐ X.docx", "type": "file"}]})
    logs = []
    r = contracts.sync_templates(drive, store, "fA,fKhongCoQuyen", log=logs.append)
    assert r["templates"] == 1
    assert any("không đọc được folder" in m for m in logs)


def test_empty_template_folders_say_so_instead_of_silent_success(tmp_path):
    """0 mẫu mà im lặng thì người vận hành tưởng S2 đã sẵn sàng."""
    store, pf, eng, g, b = make(tmp_path)
    drive = MultiFolderDrive({"fA": [], "fB": []})
    logs = []
    assert contracts.sync_templates(drive, store, "fA,fB", log=logs.append)["templates"] == 0
    assert any("chưa bỏ mẫu vào" in m or "[MAU]" in m for m in logs)


# ============ mẫu THẬT: chỗ trống dạng "………", không phải {{...}} ============

def _docx_real_style():
    """Dựng docx giống mẫu thật của legal team: chỗ trống là dãy ba chấm giữa câu,
    thông tin bên A trong bảng `nhãn | : | giá trị` có ô GỘP ở hàng đầu."""
    from docx import Document
    d = Document()
    d.add_paragraph("Số: ………/2026/HDMBHH/…….- HTC")
    d.add_paragraph("Thời gian giao hàng: …….. giờ ngày …………….;")
    d.add_paragraph("Hàng hóa là hàng mới 100%... theo yêu cầu.")   # dấu lược, KHÔNG điền
    t = d.add_table(rows=3, cols=3)
    t.rows[0].cells[0].merge(t.rows[0].cells[2]).text = "BÊN A: ….."
    t.rows[1].cells[0].text, t.rows[1].cells[1].text = "Mã số thuế", ":"
    t.rows[1].cells[2].text = "….."
    t.rows[2].cells[0].text, t.rows[2].cells[1].text = "Địa chỉ trụ sở", ":"
    t.rows[2].cells[2].text = "….."
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def test_blanks_detected_from_real_style_template():
    raw = _docx_real_style()
    bl = contracts.blanks_in_docx(raw)
    labels = [b["label"] for b in bl]
    # 2 (số HĐ) + 2 (giờ/ngày) + 1 (BÊN A) + 2 (bảng) = 7; dấu lược "100%..." KHÔNG tính
    assert len(bl) == 7, labels
    assert any("Mã số thuế" in l for l in labels)
    assert any("Địa chỉ trụ sở" in l for l in labels)
    assert any("Thời gian giao hàng" in l for l in labels)


def test_merged_cell_counted_once():
    """python-docx trả ô GỘP nhiều lần. Không lọc thì một chỗ trống đếm thành ba ⇒ lệch
    số thứ tự ⇒ MỌI giá trị phía sau rơi sai ô."""
    bl = contracts.blanks_in_docx(_docx_real_style())
    assert sum(1 for b in bl if "BÊN A" in b["label"]) == 1


def test_plain_ellipsis_is_punctuation_not_a_blank():
    """"..." trong "hàng mới 100%..." là dấu lược. Điền vào đó là làm hỏng câu."""
    from docx import Document
    d = Document()
    d.add_paragraph("Hàng hóa là hàng mới 100%... theo yêu cầu.")
    out = io.BytesIO(); d.save(out)
    assert contracts.blanks_in_docx(out.getvalue()) == []


def test_fill_uses_same_order_as_detect_and_keeps_unfilled_blanks():
    """Dò và điền phải đi CÙNG một đường. Và chỗ trống không có giá trị thì để nguyên —
    xoá âm thầm là tạo hợp đồng thiếu điều khoản mà trông như đã hoàn chỉnh."""
    from docx import Document
    raw = _docx_real_style()
    bl = contracts.blanks_in_docx(raw)
    by_label = {b["label"]: b["key"] for b in bl}
    mst = next(k for l, k in by_label.items() if "Mã số thuế" in l)
    filled = contracts.fill_docx(raw, {mst: "0108240335"})
    doc = Document(io.BytesIO(filled))
    cells = [c.text for t in doc.tables for r in t.rows for c in r.cells]
    assert "0108240335" in " ".join(cells)
    body = "\n".join(p.text for p in doc.paragraphs)
    assert "………" in body, "chỗ trống chưa điền phải còn nguyên"
    assert "100%..." in body, "dấu lược không được thay"
    assert contracts.DRAFT_MARK in body


def test_only_marked_templates_are_registered(tmp_path):
    """Trong folder mẫu còn báo giá, đề nghị thanh toán, biên bản nghiệm thu — tên đều có
    chữ "hợp đồng". Lấy bừa theo từ khoá là agent đem biên bản nghiệm thu ra soạn HĐ."""
    store, pf, eng, g, b = make(tmp_path)
    drive = MultiFolderDrive({
        "root": [{"token": "fMB", "name": "BỘ MẪU HỢP ĐỒNG_Mua bán", "type": "folder"},
                 {"token": "fHF", "name": "MẪU_Header_Footer", "type": "folder"}],
        "fMB": [{"token": "t1", "name": "2. [MAU]_Hop dong mua ban_Hapas.docx", "type": "file"},
                {"token": "t2", "name": "4. Biên bản nghiệm thu thanh lý Hợp đồng.docx",
                 "type": "file"},
                {"token": "t3", "name": "1. Bao gia_AGRI.docx", "type": "file"}],
        "fHF": [{"token": "t9", "name": "Header_Footer_HPAS.docx", "type": "file"}],
    })
    logs = []
    assert contracts.sync_templates(drive, store, "root", log=logs.append)["templates"] == 1
    assert contracts.templates(store)[0]["name"].startswith("Mua bán · ")
    assert any("bỏ qua" in m for m in logs), "phải NÓI ra file nào bị bỏ, không im lặng"


def test_review_card_shows_labels_not_field_keys(tmp_path):
    """Card ghi `cho_trong_1=015` thì người duyệt không biết đó là số hợp đồng hay tên
    bên A. Nhãn lấy từ mảnh câu quanh chỗ trống."""
    t = {"name": "Mua bán · X", "fields": [
        {"key": "cho_trong_1", "label": "Số: ___ /2026/HDMBHH"},
        {"key": "cho_trong_8", "label": "BÊN A: ….. — BÊN A:"},
        {"key": "cho_trong_9", "label": "Mã số thuế"},
        {"key": "cho_trong_20", "label": "2 — – Bên B"}]}
    v = {"cho_trong_1": "015", "cho_trong_8": "CÔNG TY ABC", "cho_trong_9": ""}
    out = flows._values_summary(t, v)
    assert "Số: ___ /2026/HDMBHH: 015" in out
    assert "BÊN A" in out and "CÔNG TY ABC" in out
    assert "cho_trong" not in out
    assert "Mã số thuế" not in out, "field rỗng không đưa vào tóm tắt"
    # 3 field chưa điền: cho_trong_9 (rỗng), cho_trong_20 (thiếu hẳn)
    assert flows._n_unfilled(t, v) == 2


def test_card_warns_about_unfilled_blanks(tmp_path):
    """Người duyệt mở file thấy dấu ……… mà không được báo trước sẽ tưởng agent làm thiếu."""
    store, pf, eng, g, b = make(tmp_path)
    gid = g.open("s2_draft", GATE, risk="medium", title="HĐ mua bán",
                 payload={"summary": "Số HĐ: 015", "n_blank": 8,
                          "file": "https://tenant/file/x"})
    card = g.render(g.get(gid))
    assert "Còn trống:** 8 chỗ" in card and "không xoá" in card
